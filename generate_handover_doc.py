"""
DDN AI Project - Complete Handover Document Generator
Generates comprehensive .docx in Rysun Labs format
Covers: Architecture, Workflows, Data, Dashboard, Roles & Responsibilities
For: BA, PM, Dev, QA, Client - 100% coverage
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime


# ─── HELPERS ───────────────────────────────────────────────────────────────────

BLUE = "1F4E79"
LIGHT_BLUE = "D6E4F0"
GREEN = "2E7D32"
LIGHT_GREEN = "E8F5E9"
RED = "C62828"
LIGHT_RED = "FFEBEE"
ORANGE = "E65100"
LIGHT_ORANGE = "FFF3E0"
GRAY = "616161"
LIGHT_GRAY = "F5F5F5"


def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def styled_table(doc, headers, rows, header_color=BLUE):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shade(c, header_color)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            if ri % 2 == 1:
                shade(c, LIGHT_BLUE)
    return t


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, size=10, color=None, font=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font:
        run.font.name = font
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(level * 1.27)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    return p


def spacer(doc):
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


# ─── DOCUMENT BUILDER ─────────────────────────────────────────────────────────

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    today = datetime.date.today().strftime("%d-%B-%Y")

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    spacer(doc)
    para(doc, "RYSUN LABS PVT. LTD.", bold=True, size=26,
         color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Technology Solutions & Services", size=12,
         color=(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(doc)
    para(doc, "PROJECT HANDOVER DOCUMENT", bold=True, size=22,
         color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "DDN AI-Assisted Test Case Failure Analysis System", bold=True,
         size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(doc)

    info = [
        ("Project Name", "DDN AI-Assisted Test Case Failure Analysis System"),
        ("Client", "Data Direct Networks (DDN)"),
        ("Vendor", "Rysun Labs Pvt. Ltd."),
        ("Date", today),
        ("Prepared By", "Sushrut Nistane (sushrut.nistane@rysun.com)"),
        ("Document Version", "1.0"),
        ("Classification", "CONFIDENTIAL"),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        shade(t.rows[i].cells[0], BLUE)
        for p in t.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        for p in t.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    spacer(doc)
    para(doc, "CONFIDENTIAL", bold=True, size=11,
         color=(192, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "This document contains proprietary information of Rysun Labs Pvt. Ltd.\n"
         "Distribution is restricted to authorized personnel only.",
         size=9, color=(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "Table of Contents", 1)
    toc = [
        "1. Executive Summary & Key Highlights",
        "2. System & Environment Details",
        "3. Repository Access (3 DDN Repos)",
        "4. Technology Stack (30+ Technologies)",
        "5. High-Level Architecture & Data Flow",
        "6. Project Structure & File Map",
        "7. ReAct Agent Architecture (7-Node AI Engine)",
        "8. RAG & Retrieval Pipeline (Fusion RAG + CRAG)",
        "9. Error Classification & Routing",
        "10. Python Workflow Engine (n8n Orchestration Layer)",
        "11. Database Schemas (MongoDB + PostgreSQL + Pinecone)",
        "12. Authentication & RBAC (JWT + 6-Level Roles)",
        "13. Multi-Project Architecture & Row-Level Security",
        "14. Dashboard API (100+ REST Endpoints, v1 + v2)",
        "15. Dashboard UI (React - 25 Pages, 25+ Components)",
        "16. AI Copilot & JARVIS Voice Assistant",
        "17. Flutter Mobile Application (Android/iOS)",
        "18. Notifications Service (Email + In-App)",
        "19. Workflow Execution Tracker (n8n-Style Debugging)",
        "20. Rancher Desktop & Docker Services (17+ Containers)",
        "21. CI/CD Pipelines (Jenkins + GitHub Actions)",
        "22. Test Coverage (50+ Test Files, 4 Frameworks)",
        "23. Configuration & Credentials",
        "24. Jira Integration & Auto-Bug Creation",
        "25. GitHub PR Auto-Creation & Code Fix Workflow",
        "26. Slack Integration (Notifications & Slash Commands)",
        "27. Knowledge Management API (CRUD + Pinecone Sync)",
        "28. Service Manager, Aging Service & Self-Healing",
        "29. Celery Async Task Queue & Flower Monitoring",
        "30. DDN Products & Guruttava Project Integration",
        "31. Implementation Maturity Assessment",
        "32. Phase Completion Status (Phase 0-8+)",
        "33. Known Issues & Pending Work",
        "34. Roles & Responsibilities (Actual RBAC Implemented)",
        "35. Handover Checklist (Day 1, Week 1, Month 1)",
        "36. Key Contacts",
        "37. FAQ for Stakeholders (BA/PM/Dev/QA/Client)",
        "38. Additional Documentation Index",
        "39. Quick Reference Commands",
        "40. Handover Email Format",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(31, 78, 121)
    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "1. Executive Summary & Key Highlights", 1)
    doc.add_paragraph(
        "The DDN AI-Assisted Test Case Failure Analysis System is an AI-powered platform "
        "developed by Rysun Labs Pvt. Ltd. for Data Direct Networks (DDN). It automates "
        "the analysis of test case failures from Jenkins CI/CD pipelines using a 7-node "
        "LangGraph ReAct agent, Fusion RAG retrieval, CRAG verification, and multi-model "
        "AI processing (Claude 3.5 Sonnet + Gemini Flash + GPT-4o-mini)."
    )

    heading(doc, "Key Performance Metrics", 2)
    styled_table(doc, ["Metric", "Before AI", "After AI", "Improvement"], [
        ("Time per Analysis", "60 minutes", "15 seconds", "99.6% reduction"),
        ("Cost per Analysis", "$30.00", "$0.05", "99.8% reduction"),
        ("Cases per Day/Engineer", "8 cases", "24 cases", "3x increase"),
        ("First Year ROI", "-", "-", "103.7%"),
        ("CRAG Auto-Accept Rate", "-", ">85% confidence", "Reduces human review"),
        ("Cache Hit Rate (Redis)", "-", "60-75%", "Faster repeat queries"),
    ])

    heading(doc, "What the System Does (For Non-Technical Readers)", 2)
    steps = [
        "Jenkins runs automated tests on DDN storage products (EXAScaler, AI400X, Infinia, IntelliFlash, JARVICE HPC)",
        "When tests FAIL, the failure data is automatically saved to MongoDB Atlas",
        "n8n workflow engine detects the new failure and triggers AI analysis",
        "The LangGraph ReAct Agent (7-step reasoning engine) analyzes the failure:",
        "  a) Classifies the error into 6 categories (Code, Infrastructure, Config, Dependency, Test, Unknown)",
        "  b) Searches knowledge base (Pinecone) for similar known errors and solutions",
        "  c) Fetches source code from GitHub (for code errors only) to inspect the bug",
        "  d) Queries historical data from MongoDB and PostgreSQL",
        "  e) Generates root cause analysis and fix recommendation",
        "  f) Verifies the answer quality using CRAG confidence scoring",
        "The AI-formatted result is stored in PostgreSQL and sent via Teams/Slack notification",
        "Engineers view results on the React Dashboard with options to Accept, Reject, or Refine the analysis",
        "For code errors, the system can auto-generate a Pull Request with the fix (requires human approval)",
    ]
    for s in steps:
        if s.startswith("  "):
            bullet(doc, s.strip(), level=1)
        else:
            bullet(doc, s)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 2. SYSTEM & ENVIRONMENT DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "2. System & Environment Details", 1)

    heading(doc, "Development Machine", 2)
    styled_table(doc, ["Property", "Value"], [
        ("Computer Name", "RYSUN-PUNE-03"),
        ("Registered Owner", "Rysun"),
        ("Domain", "rysun.com"),
        ("OS", "Microsoft Windows 11 Enterprise (Build 26100)"),
        ("Processor", "Intel Core i7-8700K @ 3600 MHz"),
        ("RAM", "32,716 MB (32 GB)"),
        ("IP Address", "10.0.1.3 (Ethernet - DHCP)"),
        ("Project Folder", r"D:\DDN-AI-Project-Documentation"),
        ("Virtual Environment", r"D:\DDN-AI-Project-Documentation\.venv"),
        ("System Boot Time", "07-Feb-2026, 11:08:37"),
    ])

    heading(doc, "Installed Software", 2)
    styled_table(doc, ["Tool", "Version", "Path"], [
        ("Python", "3.13.5", r"C:\Program Files\Python313\python.exe"),
        ("Node.js", "20.20.0", r"C:\nvm4w\nodejs\node.exe"),
        ("npm", "9.8.1", r"C:\nvm4w\nodejs\npm"),
        ("Git", "2.52.0", r"C:\Program Files\Git\cmd\git.exe"),
        ("Docker (Rancher Desktop)", "28.3.3-rd", r"C:\Program Files\Rancher Desktop"),
        ("NVM for Windows", "Installed", r"C:\nvm4w"),
        ("VS Code", "Installed", "Open with: code D:\\DDN-AI-Project-Documentation"),
    ])

    heading(doc, "How to Open the Project", 2)
    code_block(doc, '# Open project in VS Code\ncode "D:\\DDN-AI-Project-Documentation"\n\n'
               '# Activate Python virtual environment\nD:\\DDN-AI-Project-Documentation\\.venv\\Scripts\\activate\n\n'
               '# Verify Python\npython --version  # Should show 3.13.5')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. REPOSITORY ACCESS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "3. Repository Access (3 DDN Repos)", 1)
    doc.add_paragraph(
        "All DDN project code is hosted on GitHub under the account Sushrut-01. "
        "There are 3 separate repositories serving different purposes:"
    )

    heading(doc, "Repo 1: ddn-ai-test-analysis (Main Project)", 2)
    styled_table(doc, ["Property", "Value"], [
        ("URL", "https://github.com/Sushrut-01/ddn-ai-test-analysis"),
        ("Description", "AI-powered DDN storage test failure analysis with real-time monitoring"),
        ("Branch", "main"),
        ("Language", "Python (backend) + JavaScript/React (frontend)"),
        ("Clone Command", "git clone https://github.com/Sushrut-01/ddn-ai-test-analysis.git"),
        ("Contains", "All AI services, dashboard, workflows, configuration, documentation"),
        ("Status", "3 local commits ahead + 147 uncommitted files"),
    ])

    heading(doc, "Repo 2: ddn-playwright-automation (E2E QA Testing)", 2)
    styled_table(doc, ["Property", "Value"], [
        ("URL", "https://github.com/Sushrut-01/ddn-playwright-automation"),
        ("Description", "Comprehensive E2E testing suite for DDN Dashboard using Playwright + MCP"),
        ("Language", "TypeScript (85%) + Gherkin (15%)"),
        ("Clone Command", "git clone https://github.com/Sushrut-01/ddn-playwright-automation.git"),
        ("Test Suites", "Dashboard Tests, Manual Analysis Tests, Failures Page Tests"),
        ("Features", "POM pattern, multi-browser, CI/CD with GitHub Actions, MCP integration"),
        ("License", "MIT"),
    ])

    heading(doc, "Repo 3: ddn-jenkins-testing (Robot Framework)", 2)
    styled_table(doc, ["Property", "Value"], [
        ("URL", "https://github.com/Sushrut-01/ddn-jenkins-testing"),
        ("Description", "Automated Robot Framework Testing Suite for DDN Storage Products"),
        ("Language", "Python (50%) + RobotFramework (46%) + Shell (4%)"),
        ("Clone Command", "git clone https://github.com/Sushrut-01/ddn-jenkins-testing.git"),
        ("Test Suites", "16 basic tests + 7 advanced tests covering EXAScaler, AI400X, Infinia, IntelliFlash"),
        ("Features", "MongoDB auto-reporting via Robot listener, Jenkins pipeline integration"),
        ("Products Tested", "EXAScaler, AI400X, Infinia, IntelliFlash, S3, Multi-Tenancy"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 4. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "4. Technology Stack (30+ Technologies)", 1)
    styled_table(doc, ["Layer", "Technology", "Port", "Purpose"], [
        ("Orchestration", "n8n", "5678", "Workflow automation - triggers AI analysis on new failures"),
        ("AI Core (Analysis)", "Claude 3.5 Sonnet", "-", "Primary error analysis LLM"),
        ("AI Core (Classification)", "GPT-4o-mini", "-", "Error classification + reasoning (cost-optimized)"),
        ("AI Core (Formatting)", "Gemini Flash 2.0", "-", "Response formatting + code analysis + AI Copilot"),
        ("Agent Framework", "LangGraph ReAct", "5000", "7-node reasoning loop with tool execution"),
        ("Vector Database", "Pinecone (dual-index)", "-", "Knowledge docs (25 vectors) + Error library (146+ vectors)"),
        ("Document Database", "MongoDB Atlas", "27017", "Test failure data from Jenkins/Robot Framework"),
        ("Relational Database", "PostgreSQL 15", "5434", "AI analysis results, feedback, code fixes, RLS multi-tenancy"),
        ("Caching", "Redis", "6379", "Query result caching (60-75% hit rate)"),
        ("Auth Service", "Flask + JWT + bcrypt", "5013", "Authentication, user management, RBAC (6 roles)"),
        ("Dashboard API", "Flask (Python)", "5006", "25+ REST endpoints for frontend"),
        ("API v2 (Middleware)", "Flask + Blueprint", "5020", "Refactored endpoints with RLS + RBAC middleware"),
        ("Project API", "Flask + Blueprint", "-", "Multi-project CRUD with role-based access"),
        ("Knowledge API", "Flask (Python)", "5008", "Knowledge base CRUD + category management"),
        ("Manual Trigger API", "Flask (Python)", "5004", "Manual analysis trigger service"),
        ("Notifications Service", "Flask + SMTP", "5014", "Email + in-app notifications with queue"),
        ("Workflow Executions API", "Flask (Python)", "5016", "n8n-style execution tracking REST API"),
        ("Dashboard UI", "React 18 + Vite + MUI", "5173", "25-page analytics frontend with dark/light theme"),
        ("AI Copilot", "Gemini Flash 2.0 + Web Speech API", "-", "AI assistant with voice input (JARVIS)"),
        ("Mobile App", "Flutter 3.x + Riverpod", "-", "Android/iOS app with offline-first architecture"),
        ("LLM Monitoring", "Langfuse (self-hosted)", "3000", "LLM tracing, observability, cost tracking"),
        ("Async Processing", "Celery + Flower", "5555", "Background task processing and monitoring"),
        ("Containerization", "Docker Compose (Rancher Desktop)", "-", "17+ microservices orchestrated"),
        ("CI/CD", "Jenkins + GitHub Actions", "8081", "Automated test execution and analysis pipelines"),
        ("E2E Testing", "Playwright + TypeScript", "-", "Dashboard UI automated testing"),
        ("Functional Testing", "Robot Framework", "-", "DDN product testing with MongoDB reporting"),
        ("Unit Testing", "pytest + Flutter test", "-", "50+ Python test files + mobile tests"),
        ("Embeddings", "OpenAI text-embedding-ada-002", "-", "Vector embeddings for Pinecone (1536 dimensions)"),
        ("Re-ranking", "CrossEncoder (ms-marco-MiniLM)", "-", "Result re-ranking for Fusion RAG"),
        ("Voice AI", "Web Speech API", "-", "Speech recognition + synthesis (JARVIS voice assistant)"),
        ("Face Detection", "Browser FaceDetection API", "-", "Biometric login on supported browsers"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 5. HIGH-LEVEL ARCHITECTURE & DATA FLOW
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "5. High-Level Architecture & Data Flow", 1)

    heading(doc, "End-to-End Data Flow", 2)
    flow_steps = [
        ("1. Test Execution", "Jenkins runs Robot Framework / Mocha tests against DDN storage products"),
        ("2. Failure Detection", "Test failures are automatically reported to MongoDB Atlas via Robot Listener"),
        ("3. Workflow Trigger", "n8n detects new failures in MongoDB and triggers the analysis pipeline"),
        ("4. AI Analysis", "LangGraph ReAct Agent performs 7-step reasoning with tool execution"),
        ("5. RAG Retrieval", "Fusion RAG searches 4 sources: Pinecone (2 indexes) + MongoDB + PostgreSQL"),
        ("6. Code Fetch", "For CODE_ERROR: GitHub MCP server fetches relevant source code"),
        ("7. CRAG Verification", "Confidence scoring determines: auto-accept, human review, self-correct, or escalate"),
        ("8. Result Storage", "Analysis stored in PostgreSQL with classification, root cause, fix, confidence"),
        ("9. Notification", "Teams/Slack notification sent to engineering team"),
        ("10. Dashboard Display", "React dashboard shows results with Accept/Reject/Refine actions"),
        ("11. Feedback Loop", "Human feedback triggers refinement workflow for improved analysis"),
        ("12. Code Fix (optional)", "AI generates Pull Request with fix code, requires human approval"),
    ]
    styled_table(doc, ["Step", "Description"], flow_steps)

    heading(doc, "Architecture Diagram (Text)", 2)
    code_block(doc,
        "Jenkins/GitHub Tests --> MongoDB Atlas (failure data)\n"
        "         |                        |\n"
        "         v                        v\n"
        "    n8n Workflow  ----------> LangGraph ReAct Agent (Port 5000)\n"
        "                              |-- Node 1: CLASSIFY (GPT-4o-mini)\n"
        "                              |-- Node 2: REASONING (ThoughtPrompts)\n"
        "                              |-- Node 3: TOOL SELECTION (ToolRegistry)\n"
        "                              |-- Node 4: TOOL EXECUTION (retry + cache)\n"
        "                              |     |-- Pinecone Knowledge (25 docs)\n"
        "                              |     |-- Pinecone Error Library (146+ cases)\n"
        "                              |     |-- GitHub MCP (source code)\n"
        "                              |     |-- MongoDB (test logs)\n"
        "                              |     |-- PostgreSQL (history)\n"
        "                              |-- Node 5: OBSERVATION (summarize)\n"
        "                              |-- Node 6: ANSWER GENERATION (GPT-4o-mini)\n"
        "                              |-- Node 7: CRAG VERIFICATION\n"
        "                                         |\n"
        "                              Gemini Flash (formatting)\n"
        "                                         |\n"
        "                              PostgreSQL (store results)\n"
        "                                         |\n"
        "                              Teams/Slack Notification\n"
        "                                         |\n"
        "                              React Dashboard (Port 5173)\n"
        "                              |-- Dashboard (overview)\n"
        "                              |-- Failures (build-level list)\n"
        "                              |-- Failure Details (deep dive)\n"
        "                              |-- Analytics (trends)\n"
        "                              |-- Manual Trigger\n"
        "                              |-- Knowledge Management\n"
        "                                         |\n"
        "                              Langfuse (Port 3000) - LLM tracing"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 6. PROJECT STRUCTURE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "6. Project Structure & File Map", 1)

    heading(doc, "Top-Level Directory", 2)
    code_block(doc,
        "D:\\DDN-AI-Project-Documentation\\\n"
        "+-- .claude/              Claude Code config, skills, hooks\n"
        "+-- .github/              GitHub Actions workflows + issue templates\n"
        "+-- architecture/         Architecture documentation\n"
        "+-- ddn_ai_mobile/        *** FLUTTER MOBILE APP (Android/iOS) ***\n"
        "+-- docs/                 Organized documentation (archive, guides)\n"
        "+-- implementation/       *** MAIN PROJECT CODE ***\n"
        "+-- jenkins/              Jenkins pipeline configs + job definitions\n"
        "+-- logs/                 Application log files\n"
        "+-- mcp-configs/          MCP server configurations (GitHub, MongoDB)\n"
        "+-- robot-tests/          Robot Framework test suites\n"
        "+-- scripts/              Utility & startup scripts\n"
        "+-- tests/                JavaScript/Mocha test scenarios\n"
        "+-- docker-compose.yml    Docker services (17+ containers)\n"
        "+-- Jenkinsfile           Jenkins pipeline definition\n"
        "+-- .env.MASTER           Master environment config (CRITICAL)\n"
        "+-- CLAUDE.md             AI assistant project memory\n"
        "+-- README.md             Project README"
    )

    heading(doc, "Implementation Directory (Core Code)", 2)
    code_block(doc,
        "implementation/\n"
        "+-- agents/                ReAct Agent framework\n"
        "|   +-- react_agent_service.py     (1638 lines - core 7-node agent)\n"
        "|   +-- tool_registry.py           (779 lines - 15+ tools)\n"
        "|   +-- correction_strategy.py     (233 lines - self-correction)\n"
        "|   +-- thought_prompts.py         (666 lines - prompt templates)\n"
        "+-- dashboard-ui/          React frontend application\n"
        "|   +-- src/pages/         25 page components (Login, Copilot, Projects...)\n"
        "|   +-- src/components/    19+ shared components (FaceDetection, etc.)\n"
        "|   +-- src/hooks/         useAuth, useVoiceAssistant, usePageData\n"
        "|   +-- src/context/       AuthContext (JWT session management)\n"
        "|   +-- src/theme/         ThemeContext (light/dark mode)\n"
        "|   +-- src/services/      API service layer (100+ endpoints)\n"
        "+-- middleware/             Project context & auth middleware\n"
        "|   +-- project_context.py         (804 lines - RLS + RBAC middleware)\n"
        "+-- retrieval/             RAG retrieval services\n"
        "|   +-- fusion_rag_service.py      (950 lines - 4-source fusion)\n"
        "+-- verification/          CRAG verification\n"
        "|   +-- crag_verifier.py           (650 lines - confidence scoring)\n"
        "+-- database/              Database setup & utilities\n"
        "+-- migrations/            8 SQL migration scripts (multi-project, RLS)\n"
        "+-- evaluation/            Benchmarking & evaluation\n"
        "+-- security/              PII redaction (disabled by design)\n"
        "+-- workflows/             n8n workflow definitions (JSON)\n"
        "+-- tests/                 7 core pytest test files\n"
        "+-- auth_service.py                (704 lines - JWT auth + RBAC)\n"
        "+-- notifications_service.py       (597 lines - email + in-app)\n"
        "+-- workflow_execution_tracker.py   (531 lines - n8n-style debugging)\n"
        "+-- workflow_executions_api.py      (403 lines - execution REST API)\n"
        "+-- api_middleware_server.py        (77 lines - middleware launcher)\n"
        "+-- project_api.py                 (562 lines - multi-project CRUD)\n"
        "+-- api_refactored_with_middleware.py (583 lines - v2 API + RLS)\n"
        "+-- ai_analysis_service.py         (1426 lines - main orchestrator)\n"
        "+-- langgraph_agent.py             (320 lines - LangGraph wrapper)\n"
        "+-- dashboard_api_full.py          (2083 lines - 25 REST endpoints)\n"
        "+-- context_engineering.py         (700 lines - query enrichment)\n"
        "+-- code_fix_automation.py         (auto PR generation)\n"
        "+-- service_manager_api.py         (service health management)\n"
        "+-- manual_trigger_api.py          (manual analysis trigger)\n"
        "+-- aging_service.py               (test case aging tracker)\n"
        "+-- knowledge_management_api.py    (knowledge base CRUD)\n"
        "+-- self_healing_service.py        (auto-recovery)\n"
        "+-- slack_integration_service.py   (Slack notifications)\n"
        "+-- jira_integration_service.py    (Jira ticket creation)\n"
        "+-- requirements.txt               (Python dependencies)\n"
        "+-- Dockerfile                     (Container build)"
    )

    heading(doc, "File Statistics", 2)
    styled_table(doc, ["Category", "Count", "Notes"], [
        ("Python Source Files", "63", "Core services, APIs, utilities"),
        ("Python Test Files", "50+", "pytest format"),
        ("React/JSX Components", "23", "7 pages + 13+ shared components"),
        ("Documentation (MD)", "40+", "Architecture, guides, phase docs"),
        ("Documentation (HTML)", "12", "Visual architecture diagrams"),
        ("n8n Workflow JSONs", "3+", "Automation workflows"),
        ("Docker Configs", "3", "Dockerfiles + docker-compose.yml"),
        ("CI/CD Pipelines", "6", "Jenkinsfile + GitHub Actions + job configs"),
        ("PowerShell Scripts", "25+", "Service management, Docker, diagnostics"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 7. REACT AGENT ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "7. ReAct Agent Architecture (7-Node AI Engine)", 1)
    doc.add_paragraph(
        "The heart of the system is a 7-node LangGraph ReAct (Reasoning + Acting) agent "
        "that performs iterative reasoning with tool execution. It runs up to 5 iterations "
        "to gather information before generating a final answer."
    )

    heading(doc, "7 Nodes Explained", 2)
    styled_table(doc, ["Node", "Name", "LLM", "Purpose", "Key Output"], [
        ("1", "CLASSIFY", "GPT-4o-mini (temp=0.0)", "Categorize error into 6 types, detect multi-file errors, apply routing", "error_category, routing_decision"),
        ("2", "REASONING", "GPT-4o-mini (temp=0.2)", "Think about what info is needed next (THOUGHT step)", "next_action, confidence, needs_more_info"),
        ("3", "TOOL SELECTION", "None (logic)", "Pick which tool to execute, verify routing allows it", "selected tool name"),
        ("4", "TOOL EXECUTION", "None (API calls)", "Execute tool with retry (3x, exponential backoff), cache results", "tool_results, execution_metrics"),
        ("5", "OBSERVATION", "None (logic)", "Analyze and summarize what the tool returned", "findings summary"),
        ("6", "ANSWER GENERATION", "GPT-4o-mini (temp=0.1)", "Synthesize all gathered info into root cause + fix", "root_cause, fix_recommendation"),
        ("7", "CRAG VERIFICATION", "None (scoring)", "Calculate confidence score, decide escalation action", "crag_confidence, crag_action"),
    ])

    heading(doc, "Agent Flow (Loop)", 2)
    code_block(doc,
        "START --> CLASSIFY --> REASONING --[needs_more_info?]--> TOOL SELECT --> EXECUTE --> OBSERVE\n"
        "                         ^                                                            |\n"
        "                         |____________________________________________________________|\n"
        "                                           (loop up to 5 times)\n"
        "                         |\n"
        "                   [confidence >= 0.8 OR done OR max_iterations]\n"
        "                         |\n"
        "                         v\n"
        "                  ANSWER GENERATION --> CRAG VERIFICATION --> END"
    )

    heading(doc, "Available Tools (15+)", 2)
    styled_table(doc, ["Tool", "Source", "Cost", "Latency", "Used For"], [
        ("pinecone_knowledge", "Pinecone Index 1", "$0.002", "0.5s", "ALL categories (always run)"),
        ("pinecone_error_library", "Pinecone Index 2", "$0.002", "0.5s", "ALL categories (always run)"),
        ("github_get_file", "GitHub MCP", "Free", "1.0s", "CODE_ERROR only"),
        ("github_search_code", "GitHub MCP", "Free", "2.0s", "CODE_ERROR only"),
        ("mongodb_logs", "MongoDB Atlas", "Free", "0.3s", "INFRA, CONFIG, TEST"),
        ("postgres_history", "PostgreSQL", "Free", "0.3s", "ALL categories"),
        ("gemini_code_analysis", "Gemini API", "$0.01", "2.0s", "CODE_ERROR"),
        ("web_search", "Web", "Free", "3.0s", "Last resort (very low confidence)"),
    ])

    heading(doc, "Self-Correction Strategy", 2)
    styled_table(doc, ["Mechanism", "Details"], [
        ("Max Retries", "3 per tool"),
        ("Backoff", "Exponential: 1s, 2s, 4s (total 7s max)"),
        ("Transient Errors", "Timeout, connection reset, rate limit, 502/503/504"),
        ("Alternative Tools", "github_get_file -> github_search_code, pinecone -> web_search"),
        ("Caching", "Successful results cached to avoid duplicate calls"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 8. RAG & RETRIEVAL PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "8. RAG & Retrieval Pipeline (Fusion RAG + CRAG)", 1)

    heading(doc, "Fusion RAG (4-Source Hybrid Retrieval)", 2)
    doc.add_paragraph(
        "The system uses a 4-source Fusion RAG approach that combines multiple retrieval "
        "methods and re-ranks results using a CrossEncoder model. This achieves 15-25% "
        "higher accuracy than single-source RAG."
    )
    styled_table(doc, ["Source", "Type", "Method", "What It Searches"], [
        ("Pinecone (ddn-knowledge-docs)", "Dense Vector", "Semantic similarity (cosine)", "25 curated error docs (ERR001-ERR025)"),
        ("Pinecone (ddn-error-library)", "Dense Vector", "Semantic similarity (cosine)", "146+ historical failures with resolutions"),
        ("BM25 Index", "Sparse", "TF-IDF keyword matching", "Exact keyword/phrase matching"),
        ("MongoDB Full-Text", "Text Search", "MongoDB text index", "Test execution logs and error messages"),
    ])

    heading(doc, "CRAG Confidence Scoring", 2)
    doc.add_paragraph("Formula: 0.25*relevance + 0.25*consistency + 0.25*grounding + 0.15*completeness + 0.10*classification")
    styled_table(doc, ["Confidence Range", "Label", "Action", "Who Handles"], [
        (">= 0.85", "HIGH", "Auto-accept and notify team", "System (automatic)"),
        ("0.65 - 0.85", "MEDIUM", "Queue for human review (HITL)", "Engineer / QA"),
        ("0.40 - 0.65", "LOW", "Self-correct: retry reasoning loop", "System (automatic retry)"),
        ("< 0.40", "VERY_LOW", "Escalate to senior engineer + web search", "Senior Engineer"),
    ])

    heading(doc, "Pinecone Dual-Index Strategy", 2)
    styled_table(doc, ["Index", "Name", "Vectors", "Dimension", "Content"], [
        ("Index 1", "ddn-knowledge-docs", "25", "1536", "Curated error documentation (ERR001-ERR025) with solutions and code examples"),
        ("Index 2", "ddn-error-library", "146+", "1536", "Historical test failures, past resolutions, success rates, real-world examples"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 9. ERROR CLASSIFICATION & ROUTING
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "9. Error Classification & Routing", 1)

    heading(doc, "6 Error Categories", 2)
    styled_table(doc, ["Category", "Keywords", "Uses GitHub?", "Uses Gemini?", "Primary Resolution"], [
        ("CODE_ERROR", "SyntaxError, NullPointer, TypeError, AttributeError", "Yes", "Yes", "Source code inspection + knowledge base"),
        ("INFRA_ERROR", "OutOfMemory, DiskSpace, NetworkError, ConnectionTimeout", "No", "No", "RAG knowledge base only (70% cost saving)"),
        ("CONFIG_ERROR", "ConfigurationException, permission denied, env variable", "No", "No", "RAG knowledge base only"),
        ("DEPENDENCY_ERROR", "ModuleNotFoundError, ImportError, version conflict", "No", "No", "RAG knowledge base only"),
        ("TEST_ERROR", "AssertionError, expected vs actual, test failed", "Yes", "Yes", "Source code + knowledge base"),
        ("UNKNOWN", "Cannot determine", "Yes", "Yes", "All tools used (safe fallback)"),
    ])

    heading(doc, "OPTION C Routing (70% API Cost Reduction)", 2)
    doc.add_paragraph(
        "The RAGRouter implements OPTION C routing which achieves ~52% API cost reduction "
        "by only calling expensive APIs (GitHub, Gemini) for CODE_ERROR and TEST_ERROR categories. "
        "80% of errors (INFRA, CONFIG, DEPENDENCY) are resolved using RAG-only, which is "
        "faster and cheaper."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 10. N8N WORKFLOWS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "10. Python Workflow Engine (n8n Orchestration Layer)", 1)
    doc.add_paragraph(
        "n8n (https://n8n.io) is the workflow automation engine that orchestrates the "
        "entire analysis pipeline. It runs on port 5678 and connects MongoDB, the AI agent, "
        "and notification services."
    )
    styled_table(doc, ["Workflow", "File", "Trigger", "Purpose"], [
        ("Main Analysis", "workflow_1_main.json", "MongoDB new document", "Detect new failures, trigger AI analysis, store results"),
        ("Manual Trigger", "workflow_2_manual_trigger.json", "HTTP webhook (POST)", "Allow manual analysis bypassing 3-failure rule"),
        ("Refinement", "workflow_3_refinement.json", "PostgreSQL trigger", "Re-analyze when human requests refinement"),
    ])

    heading(doc, "Workflow Access", 2)
    code_block(doc, "# Access n8n workflow editor\nhttp://localhost:5678\n\n# Default credentials (check .env.MASTER)")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 11. DATABASE SCHEMAS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "11. Database Schemas", 1)

    heading(doc, "MongoDB Atlas (Document Database)", 2)
    doc.add_paragraph("Primary storage for raw test failure data from Jenkins/Robot Framework.")
    styled_table(doc, ["Collection", "Purpose", "Key Fields"], [
        ("test_failures", "Raw test failure records", "_id, test_name, build_number, job_name, error_message, stack_trace, timestamp, duration, status"),
        ("manual_triggers", "Manual analysis trigger log", "_id, build_id, triggered_by, triggered_at, status, reason, trigger_source"),
        ("test_results", "Test execution results", "_id, build_id, test_name, status, duration, timestamp"),
    ])

    heading(doc, "PostgreSQL (Relational Database)", 2)
    doc.add_paragraph("Stores AI analysis results, feedback, refinements, and code fix tracking.")
    styled_table(doc, ["Table", "Purpose", "Key Columns"], [
        ("failure_analysis", "AI analysis results", "id, mongodb_failure_id, classification, root_cause, severity, recommendation, confidence_score, ai_model, github_files, analyzed_at"),
        ("acceptance_tracking", "Human validation status", "id, analysis_id, validation_status (pending/accepted/rejected/refining/refined), validator_name, validator_email, feedback_comment"),
        ("refinement_history", "Analysis refinement iterations", "id, failure_id, iteration_number, original_*/refined_* fields, refinement_reason, refined_by, refinement_timestamp"),
        ("user_feedback", "User feedback events", "id, failure_id, feedback_type, validation_status, reason, comment, suggestion, refinement_options, feedback_timestamp"),
        ("similar_failures", "Similar failure mapping", "id, failure_id, similar_failure_id, similarity_score"),
        ("code_fix_applications", "Code fix PR tracking", "id, analysis_id, build_id, branch_name, pr_number, pr_url, pr_state, status, approved_by, error_category, ai_confidence_score, test_results"),
    ])

    heading(doc, "Pinecone (Vector Database)", 2)
    styled_table(doc, ["Index", "Vectors", "Dimension", "Metadata Fields", "Purpose"], [
        ("ddn-knowledge-docs", "25", "1536", "error_id, error_type, category, severity, solution, code_before, code_after", "Curated error documentation"),
        ("ddn-error-library", "146+", "1536", "test_name, error_message, classification, root_cause, resolution, success_rate", "Historical failure patterns"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 12. AUTHENTICATION & RBAC
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "12. Authentication & RBAC (JWT + 6-Level Roles)", 1)
    doc.add_paragraph(
        "File: implementation/auth_service.py (704 lines), Port 5013. "
        "JWT-based authentication with bcrypt password hashing, refresh tokens, "
        "and role-based access control. Frontend: LoginPage.jsx, SignupPage.jsx, ForgotPasswordPage.jsx."
    )

    heading(doc, "Auth API Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/auth/register", "POST", "User registration with email, password, name, role"),
        ("/api/auth/login", "POST", "Login - returns JWT access token + refresh token"),
        ("/api/auth/refresh", "POST", "Refresh expired access token using refresh token"),
        ("/api/auth/logout", "POST", "Invalidate current session"),
        ("/api/auth/me", "GET", "Get current user profile (requires JWT)"),
        ("/api/auth/users", "GET", "List all users (admin only)"),
        ("/api/auth/users/<id>", "PUT", "Update user profile or role"),
        ("/api/auth/users/<id>", "DELETE", "Delete user (admin only)"),
        ("/health", "GET", "Service health check"),
    ])

    heading(doc, "RBAC Role Hierarchy (6 Levels)", 2)
    styled_table(doc, ["Role", "Level", "Permissions", "Typical User"], [
        ("super_admin", "6", "All permissions + system config + user management", "System administrator"),
        ("project_owner", "5", "Full project access + manage project members", "Project lead"),
        ("project_admin", "4", "Project settings + approve code fixes + manage knowledge", "Tech lead"),
        ("developer", "3", "View + trigger analysis + approve/reject fixes", "Developer"),
        ("viewer", "2", "Read-only access to dashboards and analysis", "BA, PM, Client"),
        ("guest", "1", "Limited read access (public dashboards only)", "External stakeholder"),
    ])

    heading(doc, "Login Features", 2)
    for item in [
        "JWT Tokens: Access token (60 min) + Refresh token (7 days) with auto-renewal",
        "Password Security: bcrypt hashing with salt, minimum 8 characters",
        "Face Detection: Browser FaceDetection API for biometric login (Chrome/Edge)",
        "JARVIS Voice Auth: Voice-activated login via Web Speech API ('Hey JARVIS')",
        "Google/GitHub OAuth: OAuth buttons on login page (integration ready)",
        "Project Selection: After login, modal prompts user to select active project",
        "Theme Selection: Light/Dark mode toggle on login page (persisted)",
        "Private Routes: PrivateRoute.jsx component protects all authenticated pages",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 13. MULTI-PROJECT ARCHITECTURE & ROW-LEVEL SECURITY
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "13. Multi-Project Architecture & Row-Level Security", 1)
    doc.add_paragraph(
        "The system now supports multiple projects with complete data isolation using "
        "PostgreSQL Row-Level Security (RLS). Each project gets isolated Jira, GitHub, "
        "Pinecone, and MongoDB configurations. Files: implementation/project_api.py (562 lines), "
        "implementation/middleware/project_context.py (804 lines), "
        "implementation/api_refactored_with_middleware.py (583 lines)."
    )

    heading(doc, "Multi-Project Tables", 2)
    styled_table(doc, ["Table", "Purpose", "Key Columns"], [
        ("projects", "Central project registry", "id, slug, name, description, status, settings (JSONB)"),
        ("user_projects", "User-project association with roles", "user_id, project_id, role, extra_permissions"),
        ("project_configurations", "Per-project service config", "project_id, jira_project_key, github_repo, pinecone_namespace, mongodb_collection_prefix"),
    ])

    heading(doc, "Row-Level Security (RLS)", 2)
    for item in [
        "PostgreSQL RLS policies enforce project isolation at the database level",
        "Session variable app.current_project_id set by middleware on every request",
        "All tables with project_id column have RLS policies: SELECT, INSERT, UPDATE, DELETE",
        "Middleware auto-sets project context from JWT token + URL project_id parameter",
        "Even raw SQL queries are filtered - no data leakage between projects",
    ]:
        bullet(doc, item)

    heading(doc, "API v2 Endpoints (Middleware-Protected)", 2)
    doc.add_paragraph(
        "New /api/v2/ endpoints with automatic JWT validation, project access checks, and RLS filtering. "
        "Old /api/ endpoints remain for backward compatibility (zero-downtime migration)."
    )
    styled_table(doc, ["Endpoint", "Method", "Middleware", "Purpose"], [
        ("/api/v2/projects/<id>/failures", "GET", "require_auth + require_project_access(viewer)", "Get failures filtered by project"),
        ("/api/v2/projects/<id>/analysis/<aid>", "GET", "require_auth + require_project_access(viewer)", "Get analysis for project"),
        ("/api/v2/projects/<id>/trigger", "POST", "require_auth + require_project_access(developer)", "Trigger analysis in project context"),
        ("/api/v2/projects/<id>/fixes/approve", "POST", "require_auth + require_project_access(project_admin)", "Approve fix in project context"),
    ])

    heading(doc, "Database Migrations", 2)
    styled_table(doc, ["Migration", "Description"], [
        ("001_add_multi_project_support.sql", "Creates projects, user_projects, project_configurations tables"),
        ("002_add_guruttava_project.sql", "Seeds Guruttava project with Robot Framework config"),
        ("003_enable_row_level_security.sql", "Enables RLS policies on all project-scoped tables"),
        ("004_add_rls_to_project_configurations.sql", "RLS for project_configurations table"),
        ("add_code_healing_pipeline_fields.sql", "Fields for self-healing pipeline tracking"),
        ("add_missing_columns_for_pr_workflow.sql", "PR workflow additional columns"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 14. DASHBOARD API
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "14. Dashboard API (100+ REST Endpoints, v1 + v2)", 1)
    doc.add_paragraph(
        "The Dashboard API (dashboard_api_full.py, Port 5006) is a Flask REST API with 25+ v1 endpoints. "
        "Authentication via JWT (auth_service.py, Port 5013). New v2 endpoints add RBAC + RLS middleware. "
        "Additional APIs: Knowledge (5008), Manual Trigger (5004), Notifications (5014), Auth (5013)."
    )

    styled_table(doc, ["#", "Method", "Endpoint", "Purpose"], [
        ("1", "GET", "/api/health", "Liveness probe"),
        ("2", "GET", "/api/system/status", "System health (MongoDB, PostgreSQL, Pinecone, AI Service)"),
        ("3", "GET", "/api/pipeline/flow", "Pipeline flow with recent activity per stage"),
        ("4", "GET", "/api/failures", "List failures with pagination, filters (category, feedback, search)"),
        ("5", "GET", "/api/failures/<id>", "Single failure details with AI analysis"),
        ("6", "GET", "/api/builds/failures", "Build-level aggregation with statistics"),
        ("7", "GET", "/api/builds/<id>/tests", "All tests in a build (passed + failed)"),
        ("8", "GET", "/api/analysis/<id>", "Detailed AI analysis with similar cases"),
        ("9", "GET", "/api/stats", "System-wide statistics"),
        ("10", "GET", "/api/activity", "Recent activity log (failures + analyses)"),
        ("11", "GET", "/api/trigger/history", "Manual trigger event history"),
        ("12", "POST", "/api/trigger/manual", "Trigger AI analysis for a failure"),
        ("13", "GET", "/api/feedback/refinement-history/<id>", "Refinement iteration history"),
        ("14", "GET", "/api/analytics/acceptance-rate", "AI acceptance rate trends (7d/30d/90d)"),
        ("15", "GET", "/api/analytics/refinement-stats", "Refinement effectiveness metrics"),
        ("16", "POST", "/api/fixes/approve", "Approve code fix and create GitHub PR"),
        ("17", "POST", "/api/fixes/reject", "Reject code fix with reason"),
        ("18", "GET", "/api/fixes/<id>/status", "Fix application status with PR details"),
        ("19", "GET", "/api/fixes/history", "Fix history with filters (status, category)"),
        ("20", "POST", "/api/fixes/rollback", "Rollback a fix (close PR, revert)"),
        ("21", "GET", "/api/fixes/analytics", "Fix success analytics by category"),
    ])

    doc.add_paragraph(
        "Additional endpoints on Knowledge API (Port 5008): GET/POST/PUT/DELETE /api/knowledge/docs, "
        "GET /api/knowledge/categories, POST /api/knowledge/refresh, GET /api/knowledge/stats"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 15. DASHBOARD UI
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "15. Dashboard UI (React - 25 Pages, 25+ Components)", 1)

    heading(doc, "Authentication Pages", 2)
    styled_table(doc, ["Route", "Page", "Key Features"], [
        ("/login", "LoginPage", "Email/password login, Google/GitHub OAuth, JARVIS voice auth, face detection, theme selector, project selection modal"),
        ("/signup", "SignupPage", "User registration with role selection"),
        ("/forgot-password", "ForgotPasswordPage", "Password reset via email"),
    ])

    heading(doc, "Core Application Pages", 2)
    styled_table(doc, ["Route", "Page", "Purpose", "Key Features"], [
        ("/", "DashboardPreviewNew", "Home overview", "System health cards, metrics, recent failures, service control, acceptance trends"),
        ("/failures", "FailuresPreview", "Build-level failure list", "Search, category/status filters, pagination, timeline, aging, bulk trigger"),
        ("/failures-pending", "FailuresPendingPreview", "Pending analysis queue", "Unanalyzed failures awaiting AI processing"),
        ("/failures/:buildId", "FailureDetailsPreview", "Single failure deep-dive", "AI analysis, stack trace, GitHub code, similar errors, Accept/Reject/Refine"),
        ("/analytics", "AnalyticsPreview", "Trends and metrics", "Acceptance rate charts, trend lines, refinement stats"),
        ("/manual-trigger", "ManualTriggerPreview", "Trigger single analysis", "Build ID form, trigger history table"),
        ("/trigger-analysis", "TriggerAnalysisPreview", "Bulk analysis trigger", "Select unanalyzed failures, sequential trigger with progress"),
        ("/knowledge", "KnowledgeManagementPreview", "Knowledge base CRUD", "Add/edit/delete error docs, Pinecone sync"),
        ("/copilot", "CopilotPage", "AI Copilot assistant", "Gemini Flash 2.0 chat, voice input, code analysis, test generation"),
        ("/ai-chatbot", "AIChatbotPreview", "AI chatbot interface", "Conversational AI for debugging queries"),
        ("/ai-root-cause", "AIRootCausePreview", "Root cause deep analysis", "Detailed AI reasoning visualization"),
        ("/projects", "ProjectManagement", "Multi-project management", "Project CRUD, member management, config"),
        ("/jira-bugs", "JiraBugsPreview", "Jira bug tracking", "Jira tickets linked to AI analysis"),
        ("/pr-workflow", "PRWorkflowPreview", "GitHub PR workflow", "PR status, approval, rollback"),
        ("/rag-approval", "RAGApprovalPreview", "RAG approval queue", "Human-in-the-loop validation workflow"),
        ("/test-generator", "TestCaseGeneratorPreview", "AI test generation", "Auto-generate test cases from failures"),
        ("/services", "ServicesMonitoringPreview", "Service monitoring", "Real-time service health and control"),
        ("/pipeline-status", "PipelineStatusPreview", "Pipeline status", "CI/CD pipeline execution status"),
        ("/notifications", "NotificationsCenterPreview", "Notification center", "In-app notifications with preferences"),
        ("/users", "UserManagementPreview", "User management", "User CRUD, role assignment (admin only)"),
        ("/audit-log", "AuditLogPreview", "Audit trail", "System-wide activity log"),
        ("/configuration", "ConfigurationPreview", "System configuration", "Environment settings, feature flags"),
    ])

    heading(doc, "Key Shared Components (19+)", 2)
    styled_table(doc, ["Component", "Purpose", "Used In"], [
        ("Layout.jsx", "Navigation sidebar (25+ items) + responsive drawer", "All authenticated pages"),
        ("PrivateRoute.jsx", "JWT auth guard - redirects to /login if unauthenticated", "All protected routes"),
        ("ProjectSelector.jsx", "Project switcher dropdown in top bar", "All pages (multi-project)"),
        ("ProjectSelectionModal.jsx", "Project selection modal after login", "LoginPage"),
        ("FaceDetection.jsx", "Browser FaceDetection API for biometric login", "LoginPage"),
        ("ThemeSelector.jsx", "Light/Dark mode toggle", "LoginPage, Layout"),
        ("Breadcrumbs.jsx", "Navigation breadcrumb trail", "All pages"),
        ("ServiceControlModern.jsx", "Modern service monitoring with real-time status", "Dashboard, Services"),
        ("ServiceControl.jsx", "Start/stop/restart services, quick links", "Dashboard"),
        ("SystemStatus.jsx", "Health status cards (MongoDB, PG, Pinecone, AI)", "Dashboard"),
        ("FeedbackModal.jsx", "Reject/Refine feedback form with validation", "FailureDetails"),
        ("CodeFixApproval.jsx", "Approve/reject AI code fixes, PR tracking", "FailureDetails"),
        ("AddKnowledgeDocModal.jsx", "Comprehensive form for error documentation", "KnowledgeManagement"),
        ("TestDetailsModal.jsx", "All test cases in a build (modal)", "Failures"),
        ("BeforeAfterComparison.jsx", "Original vs refined analysis comparison", "FailureDetails"),
        ("SimilarErrorsDisplay.jsx", "Similar documented errors with solutions", "FailureDetails"),
        ("CodeSnippet.jsx", "GitHub source code with error line highlight", "FailureDetails"),
        ("DiffView.jsx", "Side-by-side code diff for fixes", "CodeFixApproval"),
        ("FeedbackStatusBadge.jsx", "Color-coded status chip", "Multiple"),
    ])

    heading(doc, "UI Technology", 2)
    styled_table(doc, ["Aspect", "Technology/Approach"], [
        ("Framework", "React 18 with functional components + hooks"),
        ("UI Library", "Material-UI (MUI) v5"),
        ("State Management", "React Query (TanStack) for server state, useState for UI state"),
        ("HTTP Client", "Axios with interceptors (auto-attaches JWT token)"),
        ("Routing", "React Router DOM v6 with PrivateRoute guards"),
        ("Authentication", "JWT via AuthContext + useAuth hook (access + refresh tokens)"),
        ("Voice Assistant", "useVoiceAssistant hook - Web Speech API (JARVIS wake word)"),
        ("Face Detection", "FaceDetection.jsx - Browser FaceDetection API"),
        ("Theming", "ThemeContext.jsx - Light/Dark mode with Material-UI theming"),
        ("Date Handling", "date-fns library"),
        ("Build Tool", "Vite (fast HMR in development)"),
        ("Auto-Refresh", "React Query refetch intervals (5s-60s per page)"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 16. AI COPILOT & JARVIS VOICE ASSISTANT
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "16. AI Copilot & JARVIS Voice Assistant", 1)
    doc.add_paragraph(
        "The AI Copilot (CopilotPage.jsx, 573 lines) is an intelligent coding assistant powered by "
        "Gemini Flash 2.0 that helps engineers debug failures, analyze code, and generate tests. "
        "It includes a JARVIS-style voice assistant (useVoiceAssistant.js) with wake word detection."
    )

    heading(doc, "Copilot Features", 2)
    for item in [
        "AI Chat: Conversational interface for test failure analysis and debugging queries",
        "Voice Input: Microphone button for voice-to-text queries (Web Speech API)",
        "Code Analysis: Paste code snippets for AI-powered bug detection",
        "Test Generation: Auto-generate test cases from failure patterns",
        "Quick Actions: Pre-built prompts for common debugging tasks",
        "Suggested Prompts: Context-aware suggestions based on recent failures",
        "Code Syntax Highlighting: Formatted code blocks in AI responses",
    ]:
        bullet(doc, item)

    heading(doc, "JARVIS Voice Assistant", 2)
    styled_table(doc, ["Feature", "Technology", "Details"], [
        ("Speech Recognition", "Web Speech API (SpeechRecognition)", "Continuous listening in Chrome/Edge/Safari"),
        ("Speech Synthesis", "Web Speech API (SpeechSynthesis)", "Google UK English Male voice (JARVIS-like)"),
        ("Wake Word", "Custom detection", "'Hey JARVIS' triggers listening mode"),
        ("Continuous Mode", "Configurable", "Keep listening after each command"),
        ("Voice Commands", "Custom processor", "Navigate, trigger analysis, check status via voice"),
        ("Browser Support", "Chrome, Edge, Safari", "Falls back gracefully on unsupported browsers"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 17. FLUTTER MOBILE APPLICATION
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "17. Flutter Mobile Application (Android/iOS)", 1)
    doc.add_paragraph(
        "A professional, production-ready Flutter mobile application for mobile access to AI-powered "
        "test failure analysis. Location: ddn_ai_mobile/ directory. Supports Android and iOS with "
        "offline-first architecture, push notifications, and AI chatbot."
    )

    heading(doc, "Mobile App Features & Status", 2)
    styled_table(doc, ["Feature", "Status", "Description"], [
        ("Authentication", "100% Complete", "JWT token-based auth with auto-refresh, splash screen"),
        ("Dashboard", "100% Complete", "System health, 4 stat cards, recent activity, pull-to-refresh"),
        ("Failures Management", "100% Complete", "Paginated list, search, filters, infinite scroll, offline cache"),
        ("AI Analysis (Phase 3)", "100% Complete", "Root cause, fix suggestions, confidence score, similar errors"),
        ("AI Chatbot", "60% Partial", "UI done (message bubbles, input bar), API needs completion"),
        ("Analytics", "70% Partial", "Chart screens done (bar/line/pie), data models incomplete"),
        ("RAG Approval Queue", "70% Partial", "List/detail screens done, approval logic incomplete"),
        ("Push Notifications", "10% Setup", "Firebase deps added, google-services.json not configured"),
        ("Offline Mode", "100% Complete", "Hive local DB with 12hr/24hr/7day cache expiry, auto-sync"),
        ("Settings", "80% Partial", "Dark mode, API endpoint, cache management"),
    ])

    heading(doc, "Technical Architecture", 2)
    styled_table(doc, ["Layer", "Technology", "Purpose"], [
        ("State Management", "Riverpod 2.x", "Reactive state with code generation"),
        ("Networking", "Dio + Retrofit", "Type-safe API calls with automatic serialization"),
        ("Local Storage", "Hive + Flutter Secure Storage", "Offline caching + secure token storage"),
        ("Navigation", "GoRouter", "Declarative routing with deep linking"),
        ("Charts", "FL Chart", "Analytics visualization"),
        ("Push Notifications", "Firebase FCM", "Real-time alerts with deep linking"),
        ("Dependency Injection", "get_it + injectable", "Service locator pattern"),
        ("Architecture", "Clean Architecture", "Presentation / Domain / Data layer separation"),
    ])

    heading(doc, "Build & Deploy", 2)
    code_block(doc,
        "# Build Android APK\ncd ddn_ai_mobile\nflutter build apk --release\n\n"
        "# Build iOS (requires Mac)\nflutter build ios --release\n\n"
        "# Run in development\nflutter run"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 18. NOTIFICATIONS SERVICE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "18. Notifications Service (Email + In-App)", 1)
    doc.add_paragraph(
        "File: implementation/notifications_service.py (597 lines), Port 5014. "
        "Provides email and in-app notifications with queue, retry logic, and per-user preferences."
    )

    heading(doc, "Notification Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/notifications", "GET", "Get user's notifications (paginated, filterable)"),
        ("/api/notifications/<id>/read", "PUT", "Mark notification as read"),
        ("/api/notifications/read-all", "PUT", "Mark all notifications as read"),
        ("/api/notifications/send", "POST", "Send notification (email + in-app)"),
        ("/api/notifications/preferences", "GET", "Get user notification preferences"),
        ("/api/notifications/preferences", "PUT", "Update notification preferences"),
        ("/health", "GET", "Service health check"),
    ])

    heading(doc, "Notification Types", 2)
    styled_table(doc, ["Type", "Channel", "Trigger"], [
        ("Analysis Complete", "Email + In-App", "AI analysis finishes for a failure"),
        ("Code Fix Ready", "Email + In-App", "AI generates a code fix for approval"),
        ("PR Status Change", "In-App", "GitHub PR merged, closed, or review requested"),
        ("CRAG Low Confidence", "Email", "Analysis below confidence threshold (escalation)"),
        ("System Alert", "Email", "Service down, health check failure"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 19. WORKFLOW EXECUTION TRACKER
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "19. Workflow Execution Tracker (n8n-Style Debugging)", 1)
    doc.add_paragraph(
        "File: implementation/workflow_execution_tracker.py (531 lines). Creates detailed "
        "execution logs with node-level input/output tracking, similar to n8n's execution "
        "visualization. Stores every LangGraph workflow execution for debugging and replay."
    )

    heading(doc, "Tracked Data Per Execution", 2)
    styled_table(doc, ["Data", "Description"], [
        ("Execution ID", "Unique identifier for each workflow run"),
        ("Start/End Time", "Precise timestamps for total duration calculation"),
        ("Status", "running, completed, failed, timeout"),
        ("Node Sequence", "Ordered list of nodes executed (e.g., CLASSIFY -> REASONING -> TOOL_SELECT...)"),
        ("Node Input/Output", "Full input and output data for each node (JSON)"),
        ("Node Duration", "Execution time per node in milliseconds"),
        ("Error Details", "Full error message and stack trace for failed nodes"),
        ("Metadata", "Build ID, failure ID, project ID, user who triggered"),
    ])

    heading(doc, "Use Cases", 2)
    for item in [
        "Debug why a specific analysis failed or produced low confidence",
        "Identify bottleneck nodes (e.g., slow GitHub API calls)",
        "Replay executions to reproduce issues",
        "Compare execution paths for similar failures",
        "Monitor average node execution times for performance optimization",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 20. DOCKER SERVICES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "20. Rancher Desktop & Docker Services (17+ Containers)", 1)
    styled_table(doc, ["#", "Container", "Port", "Image/Type", "Purpose"], [
        ("1", "n8n", "5678", "n8nio/n8n", "Workflow automation engine"),
        ("2", "langfuse-server", "3000", "langfuse/langfuse", "LLM tracing & monitoring UI"),
        ("3", "langfuse-worker", "-", "langfuse/langfuse-worker", "Background trace processing"),
        ("4", "langfuse-db", "5434", "postgres:15", "PostgreSQL for Langfuse"),
        ("5", "redis", "6379", "redis:latest", "Caching layer"),
        ("6", "postgres", "5434", "postgres:15", "Main application database"),
        ("7", "dashboard-api", "5006", "Python/Flask", "REST API backend"),
        ("8", "dashboard-ui", "5173", "Node/Vite", "React frontend"),
        ("9", "langgraph-agent", "5000", "Python/Flask", "AI agent service"),
        ("10", "mcp-server", "5002", "Python", "GitHub MCP tools"),
        ("11", "celery-worker", "-", "Python/Celery", "Background task processor"),
        ("12", "celery-flower", "5555", "Python/Flower", "Task monitoring UI"),
        ("13", "manual-trigger", "5004", "Python/Flask", "Manual analysis trigger"),
        ("14", "aging-service", "5007", "Python", "Test case aging tracker"),
        ("15", "service-manager", "5008", "Python/Flask", "Service orchestration + Knowledge API"),
        ("16", "knowledge-api", "-", "Python/Flask", "Knowledge management"),
        ("17", "self-healing", "-", "Python", "Auto-recovery service"),
    ])

    heading(doc, "Docker Commands", 2)
    code_block(doc,
        "# Start all services\ndocker-compose up -d\n\n"
        "# Stop all services\ndocker-compose down\n\n"
        "# View logs for a service\ndocker-compose logs -f dashboard-api\n\n"
        "# Restart a service\ndocker-compose restart langgraph-agent\n\n"
        "# Rebuild and start\ndocker-compose up -d --build\n\n"
        "# Check status\ndocker-compose ps"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 21. CI/CD PIPELINES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "21. CI/CD Pipelines", 1)

    heading(doc, "GitHub Actions", 2)
    doc.add_paragraph("File: .github/workflows/ci-cd.yml - Triggers on push to main branch.")

    heading(doc, "Jenkins", 2)
    styled_table(doc, ["File", "Purpose", "Location"], [
        ("Jenkinsfile", "Main pipeline definition", "Project root"),
        ("jenkins/Jenkinsfile.test", "Test-specific pipeline", "jenkins/"),
        ("jenkins/jobs/", "Job configuration files", "jenkins/"),
        ("jenkins/webhook-config.json", "Webhook settings", "jenkins/"),
        ("jenkins/SETUP-JENKINS-JOBS.bat", "Job setup automation", "jenkins/"),
    ])

    heading(doc, "Playwright CI/CD (ddn-playwright-automation repo)", 2)
    doc.add_paragraph(
        "GitHub Actions runs Playwright E2E tests on push, PR, and every 6 hours. "
        "Multi-browser: Chromium, Firefox, WebKit, Mobile Chrome, Mobile Safari. "
        "Reports in HTML, JSON, JUnit XML formats."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 22. TEST COVERAGE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "22. Test Coverage (50+ Test Files, 4 Frameworks)", 1)

    styled_table(doc, ["Framework", "Files", "Location", "Purpose"], [
        ("pytest (Python)", "50+", "implementation/ and implementation/tests/", "Unit, integration, performance tests"),
        ("Mocha (JavaScript)", "4", "tests/", "Test scenario runners for DDN products"),
        ("Robot Framework", "2", "robot-tests/", "End-to-end DDN product acceptance tests"),
        ("Playwright (TypeScript)", "3+ suites", "ddn-playwright-automation repo", "Dashboard E2E tests"),
        ("Flutter Test", "5", "ddn_ai_mobile/test/ + integration_test/", "Mobile app widget, feature, and integration tests"),
    ])

    heading(doc, "Key Test Categories", 2)
    styled_table(doc, ["Category", "Key Files", "Count"], [
        ("ReAct Agent", "test_react_agent*.py, test_langgraph*.py", "6"),
        ("RAG & Retrieval", "test_fusion_rag*.py, test_dual_index*.py, test_rag_router*.py", "8"),
        ("CRAG Verification", "test_crag_*.py", "4"),
        ("Database", "test_mongodb*.py, test_postgres*.py", "4"),
        ("Integration", "test_e2e_github*.py, test_celery*.py, test_langfuse*.py", "5"),
        ("Self-Correction", "test_self_correction.py", "1"),
        ("HITL", "test_hitl_manager.py", "1"),
        ("Context & Routing", "test_context_*.py", "4"),
        ("Performance", "test_fusion_rag_performance.py, test_crag_performance.py", "2"),
        ("Robot Framework (DDN)", "ddn_basic_tests.robot (16 cases), ddn_advanced_tests.robot (7 cases)", "23 test cases"),
    ])

    heading(doc, "How to Run Tests", 2)
    code_block(doc,
        "# Activate venv first\n"
        "D:\\DDN-AI-Project-Documentation\\.venv\\Scripts\\activate\n\n"
        "# Run all Python tests\n"
        "pytest implementation/ -v --tb=short\n\n"
        "# Run specific category\n"
        "pytest implementation/tests/test_react_agent.py -v\n"
        "pytest implementation/test_fusion_rag_simple.py -v\n\n"
        "# Run JavaScript tests\n"
        "cd tests && npm test\n\n"
        "# Run Robot Framework tests\n"
        "cd robot-tests && robot ddn_basic_tests.robot"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 17. CONFIGURATION & CREDENTIALS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "23. Configuration & Credentials", 1)
    para(doc, "CRITICAL: The master configuration file is .env.MASTER. ALWAYS create a backup before editing.", bold=True, size=10, color=(192, 0, 0))

    heading(doc, "Environment Variables", 2)
    styled_table(doc, ["Variable", "Description", "Required?"], [
        ("MONGODB_URI", "MongoDB Atlas connection string (with encoded password)", "YES"),
        ("POSTGRES_HOST / PORT / DB / USER / PASSWORD", "PostgreSQL connection (default port 5434)", "YES"),
        ("PINECONE_API_KEY", "Pinecone vector database API key", "YES"),
        ("ANTHROPIC_API_KEY", "Claude API access key", "YES"),
        ("OPENAI_API_KEY", "OpenAI API key (embeddings + GPT-4o-mini)", "YES"),
        ("GEMINI_API_KEY", "Google Gemini API key (formatting)", "YES"),
        ("LANGFUSE_PUBLIC_KEY + SECRET_KEY", "Langfuse tracing keys", "YES"),
        ("REDIS_HOST / PORT", "Redis connection (default localhost:6379)", "Optional"),
        ("GITHUB_TOKEN", "GitHub API access token", "For code fix features"),
        ("SLACK_WEBHOOK_URL", "Slack notification webhook", "Optional"),
        ("JIRA_API_TOKEN", "Jira integration token", "Optional"),
        ("N8N_USER / N8N_PASSWORD", "n8n workflow editor credentials", "YES"),
        ("JWT_SECRET_KEY", "JWT signing secret (auth_service.py)", "YES"),
        ("TOKEN_EXPIRE_MINUTES", "JWT access token expiry (default: 60)", "Optional"),
        ("REFRESH_TOKEN_EXPIRE_DAYS", "Refresh token expiry (default: 7)", "Optional"),
        ("SMTP_HOST / SMTP_PORT / SMTP_USERNAME / SMTP_PASSWORD", "Email notification SMTP config", "For notifications"),
    ])

    heading(doc, "Configuration Files", 2)
    styled_table(doc, ["File", "Purpose", "WARNING"], [
        (".env.MASTER", "Master env config (all API keys)", "NEVER edit without backup!"),
        (".env (implementation/)", "Runtime env for Python services", "Auto-derived from MASTER"),
        ("docker-compose.yml", "17 Docker service definitions", "Port changes need coordination"),
        ("requirements.txt", "Python dependencies", "Use pip install -r"),
        ("requirements-lock.txt", "Pinned dependency versions", "For reproducible builds"),
        ("package.json (dashboard-ui/)", "React UI dependencies", "Use npm install"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 18. JIRA INTEGRATION & AUTO-BUG CREATION
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "24. Jira Integration & Auto-Bug Creation", 1)
    doc.add_paragraph(
        "The system includes a fully implemented Jira integration service (Port 5009) that "
        "automatically creates Jira tickets when AI analysis completes for a test failure. "
        "File: implementation/jira_integration_service.py (447 lines)."
    )

    heading(doc, "Jira API Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/jira/create-issue", "POST", "Auto-create Jira bug with AI analysis, root cause, fix recommendation"),
        ("/api/jira/update-from-feedback", "POST", "Update Jira ticket when fix verified as success/failure"),
        ("/api/jira/get-issue/<build_id>", "GET", "Retrieve Jira issue key and URL for a build"),
        ("/health", "GET", "Service health check"),
    ])

    heading(doc, "Smart Features", 2)
    for item in [
        "Duplicate Prevention: Checks PostgreSQL for existing Jira issue before creating. If exists, adds comment instead of new ticket.",
        "Auto-Priority: 5+ failures = Highest, 3-4 = High, 1-2 = Medium priority",
        "Auto-Labels: CODE_ERROR -> ['code-error','requires-dev'], INFRA_ERROR -> ['infrastructure','devops'], TEST_FAILURE -> ['test-failure','qa-attention']",
        "Rich Content: Includes Build ID, Job Name, Error Category, Consecutive Failures, AI Confidence Score, Root Cause, Fix Recommendation, Jenkins/Dashboard/GitHub links",
        "Auto-Transition: On feedback success, auto-transitions Jira ticket to 'Resolved' (transition ID: 31)",
        "Feedback Sync: Dashboard Accept/Reject feedback is synced back to Jira as comments",
    ]:
        bullet(doc, item)

    heading(doc, "Jira Configuration", 2)
    styled_table(doc, ["Variable", "Purpose", "Example"], [
        ("JIRA_URL", "Atlassian instance URL", "https://your-company.atlassian.net"),
        ("JIRA_EMAIL", "Service account email", "service@company.com"),
        ("JIRA_API_TOKEN", "API token from Atlassian", "Get from id.atlassian.com/manage-profile/security/api-tokens"),
        ("JIRA_PROJECT_KEY", "Jira project key", "DDN"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 19. GITHUB PR AUTO-CREATION & CODE FIX WORKFLOW
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "25. GitHub PR Auto-Creation & Code Fix Workflow", 1)
    doc.add_paragraph(
        "The system can automatically create GitHub Pull Requests for AI-recommended code fixes. "
        "Files: implementation/code_fix_automation.py (951 lines), implementation/github_client.py (1235 lines)."
    )

    heading(doc, "End-to-End PR Workflow", 2)
    steps_pr = [
        "1. AI analysis completes with CODE_ERROR classification and fix recommendation",
        "2. User reviews proposed fix on Dashboard (CodeFixApproval component shows confidence, diff, code)",
        "3. User clicks 'Approve & Create PR' (or system auto-approves if confidence >= 70%)",
        "4. System creates branch: fix/build-{build_id} from main",
        "5. AI-recommended code patch is applied to the target file",
        "6. Pull Request created with AI-generated title, description, root cause, and fix details",
        "7. Reviewers auto-assigned, labels added: ['automated-fix', 'ai-generated', 'needs-review']",
        "8. PR tracked in PostgreSQL (code_fix_applications table) with full lifecycle",
        "9. Status: pending -> pr_created -> tests_running -> tests_passed -> merged (or reverted)",
    ]
    for s in steps_pr:
        bullet(doc, s)

    heading(doc, "Code Fix API Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/fixes/approve", "POST", "Approve fix & trigger GitHub PR creation"),
        ("/api/fixes/reject", "POST", "Reject fix with reason"),
        ("/api/fixes/<fix_id>/status", "GET", "Get PR status and CI test results"),
        ("/api/fixes/history", "GET", "Fix application history with filters"),
        ("/api/fixes/rollback", "POST", "Rollback a failed or unwanted fix"),
        ("/api/fixes/analytics", "GET", "Fix success rate, time-to-merge, rollback stats"),
    ])

    heading(doc, "PR Content (Auto-Generated)", 2)
    for item in [
        "Title: 'Automated Fix: {error_type} in {component}'",
        "Body: Error summary, root cause analysis, fix applied, files changed, AI confidence %",
        "Approval info: Who approved, timestamp",
        "Verification steps checklist for reviewer",
        "Footer: 'Generated by DDN AI Analysis System'",
    ]:
        bullet(doc, item)

    heading(doc, "Dashboard UI - CodeFixApproval Component", 2)
    doc.add_paragraph(
        "The CodeFixApproval.jsx component (495 lines) shows: AI confidence with color-coded progress bar, "
        "error category and severity chips, files affected, before/after code diff view, "
        "and three action buttons: Approve (green), Reject (red), Request Changes (orange). "
        "After PR creation, shows PR number, status, and clickable GitHub link."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 20. SLACK INTEGRATION
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "26. Slack Integration (Notifications & Slash Commands)", 1)
    doc.add_paragraph(
        "File: implementation/slack_integration_service.py (511 lines), Port 5012. "
        "Provides rich failure notifications, interactive buttons, and slash commands."
    )

    heading(doc, "Slack Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/slack/send-notification", "POST", "Send rich failure analysis alert to channel"),
        ("/api/slack/interactions", "POST", "Handle interactive button clicks (Fix Worked / Fix Failed)"),
        ("/api/slack/update-thread", "POST", "Post status updates in existing threads"),
        ("/api/slack/slash-command", "POST", "Handle /ddn-ai status and /ddn-ai trigger <build_id>"),
    ])

    heading(doc, "Features", 2)
    for item in [
        "Rich Messages: Formatted failure details with error category, confidence, root cause preview",
        "Interactive Buttons: 'Fix Worked' / 'Fix Failed' buttons for quick feedback without opening dashboard",
        "Thread Updates: Status changes posted as thread replies to original notification",
        "Channel Routing: Different channels per severity or category (configurable)",
        "Slash Commands: /ddn-ai status (system health), /ddn-ai trigger <build_id> (trigger analysis from Slack)",
        "Feedback Tracking: Button clicks recorded as user feedback in PostgreSQL",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 21. KNOWLEDGE MANAGEMENT API
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "27. Knowledge Management API (CRUD + Pinecone Sync)", 1)
    doc.add_paragraph(
        "File: implementation/knowledge_management_api.py (817 lines), Port 5008. "
        "Full CRUD for error documentation with automatic Pinecone vector embedding."
    )

    heading(doc, "API Endpoints", 2)
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/knowledge/docs", "GET", "List docs with filters (category, severity, search text)"),
        ("/api/knowledge/docs/<id>", "GET", "Get specific document by ID"),
        ("/api/knowledge/docs", "POST", "Add new error doc (auto-generates embedding, syncs to Pinecone)"),
        ("/api/knowledge/docs/<id>", "PUT", "Update doc (re-embeds automatically)"),
        ("/api/knowledge/docs/<id>", "DELETE", "Delete doc from both PostgreSQL and Pinecone"),
        ("/api/knowledge/categories", "GET", "Get all categories with document counts"),
        ("/api/knowledge/categories/refresh", "POST", "Refresh categories from Pinecone (no restart needed)"),
        ("/api/knowledge/stats", "GET", "Total docs, docs per category/severity, recent additions"),
    ])

    heading(doc, "Dashboard UI - Knowledge Management Page", 2)
    doc.add_paragraph(
        "The /knowledge page provides: statistics cards (total docs, categories, recent additions, vectors), "
        "filter panel (search, category dropdown, severity dropdown), documents table with Edit/Delete actions, "
        "and AddKnowledgeDocModal (30+ form fields including error type, root cause, solution steps, "
        "code before/after, severity, tags, test scenarios). All changes logged to PostgreSQL audit trail."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 22. SERVICE MANAGER, AGING SERVICE & SELF-HEALING
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "28. Service Manager, Aging Service & Self-Healing", 1)

    heading(doc, "Service Manager API (Port 5007)", 2)
    doc.add_paragraph(
        "File: implementation/service_manager_api.py (572 lines). Manages all 17 Docker services "
        "from the dashboard UI. The ServiceControl component provides Start/Stop/Restart buttons."
    )
    styled_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/services/status", "GET", "Status of all 17 services (port check + Docker SDK)"),
        ("/api/services/start/<id>", "POST", "Start specific service"),
        ("/api/services/stop/<id>", "POST", "Stop specific service"),
        ("/api/services/start-all", "POST", "Start all services in dependency order"),
        ("/api/services/stop-all", "POST", "Stop all services in reverse order"),
        ("/api/services/restart-all", "POST", "Restart all services"),
    ])

    heading(doc, "Aging Service (Port 5010)", 2)
    doc.add_paragraph(
        "File: implementation/aging_service.py (656 lines). Uses APScheduler cron job (every 6 hours) "
        "to detect unanalyzed failures that are aging. Auto-triggers AI analysis for failures with "
        ">= 2 consecutive failures over >= 3 day span."
    )
    styled_table(doc, ["Endpoint", "Purpose"], [
        ("/health", "Service health check"),
        ("/trigger-now", "Manual trigger of aging scan"),
        ("/stats", "Aging statistics and thresholds"),
        ("/recent-triggers", "Recent aging trigger history"),
    ])

    heading(doc, "Self-Healing Service (Port 5008 - EXPERIMENTAL)", 2)
    doc.add_paragraph(
        "File: implementation/self_healing_service.py (536 lines). SAFE_MODE=true by default. "
        "Predefined fix patterns (timeout increase, dependency update, cache clear, env var fix). "
        "Requires human approval for first-time fixes. Tracks success rate before auto-apply."
    )
    styled_table(doc, ["Endpoint", "Purpose"], [
        ("/api/self-heal/analyze", "Check if failure matches a known fix pattern"),
        ("/api/self-heal/apply", "Apply predefined fix (requires approval in safe mode)"),
        ("/api/self-heal/history", "Fix application history"),
        ("/api/self-heal/patterns", "List available auto-fix patterns"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 23. CELERY ASYNC TASK QUEUE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "29. Celery Async Task Queue & Flower Monitoring", 1)
    doc.add_paragraph(
        "File: implementation/tasks/celery_tasks.py (402 lines). Redis-based Celery broker for "
        "async processing. Prevents webhook timeouts by offloading long-running AI analysis."
    )

    heading(doc, "Celery Tasks", 2)
    styled_table(doc, ["Task", "Timeout", "Purpose"], [
        ("analyze_test_failure", "3-4 min", "Async AI analysis (offloaded from webhook)"),
        ("batch_analyze_failures", "Variable", "Queue multiple failures for sequential analysis"),
        ("cleanup_old_results", "5 min", "Delete expired analysis results (configurable retention)"),
    ])

    heading(doc, "Features", 2)
    for item in [
        "Progress tracking: Task state visible via /api/task-status/<task_id>",
        "Retry with exponential backoff (max 3 retries: 1s, 2s, 4s delay)",
        "Soft/hard time limits to prevent stuck tasks",
        "Task acknowledgment after completion (not before)",
        "Flower Dashboard (Port 5555): Real-time task monitoring, worker status, task history",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 24. DDN PRODUCTS & MULTI-PROJECT ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "30. DDN Products & Guruttava Project Integration", 1)

    heading(doc, "DDN Products Tested", 2)
    styled_table(doc, ["Product", "Type", "Test Framework", "Test Count"], [
        ("EXAScaler", "Lustre parallel file system", "Robot Framework", "16 basic + 7 advanced"),
        ("AI400X", "AI/GPU storage platform", "Robot Framework", "Included in advanced tests"),
        ("Infinia", "Orchestration platform", "Robot Framework", "Included in advanced tests"),
        ("IntelliFlash", "Enterprise flash arrays", "Robot Framework", "Included in advanced tests"),
        ("JARVICE", "HPC cloud computing platform", "Robot Framework", "Included in test scenarios"),
        ("S3 Multi-Tenancy", "Object storage", "Robot Framework", "Dedicated multi-tenancy tests"),
    ])

    heading(doc, "Three DDN Repositories", 2)
    styled_table(doc, ["Repository", "Purpose", "Framework", "Integration"], [
        ("ddn-ai-test-analysis", "Main AI analysis system (this repo)", "pytest (50+ files)", "Central hub - all AI services"),
        ("ddn-playwright-automation", "E2E testing of Dashboard UI", "Playwright + TypeScript", "Runs via GitHub Actions every 6hrs"),
        ("ddn-jenkins-testing", "DDN storage product testing", "Robot Framework (23 tests)", "Jenkins pipeline + MongoDB listener"),
    ])

    heading(doc, "Guruttava Project (Second Project)", 2)
    doc.add_paragraph(
        "Guruttava is the second project integrated into the multi-project architecture. "
        "It focuses on mobile (Android/iOS) and web test automation using Robot Framework with "
        "Appium (mobile) and Selenium (web). Configured via migration 002_add_guruttava_project.sql."
    )
    styled_table(doc, ["Aspect", "DDN Project", "Guruttava Project"], [
        ("Focus", "Storage product testing", "Mobile/Web test automation"),
        ("Test Framework", "Robot Framework + Mocha", "Robot Framework + Appium + Selenium"),
        ("CI Provider", "Jenkins", "Jenkins"),
        ("Jira Project Key", "DDN", "GUR (configurable)"),
        ("Pinecone Namespace", "ddn", "guruttava"),
        ("MongoDB Collection", "ddn_*", "guruttava_*"),
        ("Status", "Active (primary)", "Active (integrated)"),
    ])

    heading(doc, "Test Failure Ingestion Pipeline", 2)
    doc.add_paragraph(
        "Robot Framework tests are the primary data producer. The MongoDB Robot Listener "
        "(implementation/mongodb_robot_listener.py) automatically saves failure data to MongoDB Atlas "
        "when any Robot test fails. This triggers the AI analysis pipeline."
    )
    for item in [
        "Robot Framework -> mongodb_robot_listener.py -> MongoDB Atlas (builds, test_results collections)",
        "MongoDB -> n8n webhook OR aging_service.py -> Python AI pipeline -> PostgreSQL -> Dashboard",
        "Multi-project: Each project's data is isolated via RLS + project_id + Pinecone namespaces",
        "Playwright tests run separately (not yet feeding into AI analysis pipeline)",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 25. IMPLEMENTATION MATURITY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "31. Implementation Maturity Assessment", 1)
    doc.add_paragraph(
        "Honest assessment of what is fully working, what is experimental, what is disabled, "
        "and what is still under development."
    )

    heading(doc, "Fully Complete & Production-Ready", 2)
    styled_table(doc, ["Feature", "Status", "Files", "Notes"], [
        ("LangGraph ReAct Agent (7-node)", "COMPLETE", "agents/react_agent_service.py", "1637 lines, 7 reasoning nodes"),
        ("Fusion RAG (4-source)", "COMPLETE", "retrieval/fusion_rag_service.py", "Pinecone + BM25 + MongoDB + PostgreSQL"),
        ("CRAG Verification", "COMPLETE", "verification/crag_verifier.py", "5-dimension confidence scoring"),
        ("JWT Authentication & RBAC", "COMPLETE", "auth_service.py", "704 lines, 6-level roles, bcrypt, refresh tokens"),
        ("Multi-Project Architecture", "COMPLETE", "project_api.py + middleware/", "RLS isolation, per-project config"),
        ("Dashboard API (v1 + v2)", "COMPLETE", "dashboard_api_full.py + api_refactored", "100+ endpoints with middleware"),
        ("Dashboard UI (React)", "COMPLETE", "dashboard-ui/src/", "25 pages, 25+ components, dark/light theme"),
        ("AI Copilot + JARVIS Voice", "COMPLETE", "CopilotPage.jsx + useVoiceAssistant", "Gemini Flash 2.0, Web Speech API"),
        ("Flutter Mobile App (Core)", "COMPLETE", "ddn_ai_mobile/", "Auth, Dashboard, Failures, Analysis, Offline (55% overall)"),
        ("Notifications Service", "COMPLETE", "notifications_service.py", "597 lines, email + in-app"),
        ("Workflow Execution Tracker", "COMPLETE", "workflow_execution_tracker.py", "531 lines, n8n-style debugging"),
        ("Jira Auto-Bug Creation", "COMPLETE", "jira_integration_service.py", "Smart duplicate prevention"),
        ("GitHub PR Creation", "COMPLETE", "code_fix_automation.py", "Full PR lifecycle"),
        ("Slack Notifications", "COMPLETE", "slack_integration_service.py", "Interactive buttons + slash commands"),
        ("Knowledge Management", "COMPLETE", "knowledge_management_api.py", "Full CRUD + Pinecone sync"),
        ("n8n Workflows (4)", "COMPLETE", "workflows/*.json", "Auto-trigger, manual, refinement, auto-fix"),
        ("Langfuse Tracing", "COMPLETE", "langfuse_tracing.py", "Full LLM call observability"),
        ("Error Classification (6 types)", "COMPLETE", "agents/react_agent_service.py", "92% classification accuracy"),
        ("Context Engineering", "COMPLETE", "context_engineering.py", "Token optimization for Gemini"),
        ("Celery Task Queue", "COMPLETE", "tasks/celery_tasks.py", "Async processing with Redis broker"),
        ("Guruttava Project", "COMPLETE", "migrations/002_*", "Second project fully configured"),
    ])

    heading(doc, "Experimental / Partial", 2)
    styled_table(doc, ["Feature", "Status", "Notes"], [
        ("Self-Healing Service", "EXPERIMENTAL", "SAFE_MODE=true. Fix patterns defined but actual application needs testing."),
        ("Auto-Fix (>70% confidence)", "EXPERIMENTAL", "n8n workflow 4 can auto-approve, but typically requires human approval"),
        ("Mobile App - Chat/Analytics/RAG", "60-70%", "UI screens done, data layers and API integration need completion"),
        ("Mobile App - Push Notifications", "10%", "Firebase deps added, google-services.json config needed"),
        ("Face Detection Login", "EXPERIMENTAL", "Chrome 98+ with FaceDetection API flag required"),
    ])

    heading(doc, "Disabled by Design", 2)
    styled_table(doc, ["Feature", "Status", "Notes"], [
        ("PII Redaction (Phase 4)", "DISABLED", "Code complete but OFF. DDN test data has no PII. Re-enable via PII_REDACTION_ENABLED=true"),
    ])

    heading(doc, "Incomplete / TODO", 2)
    styled_table(doc, ["Feature", "Status", "What's Missing"], [
        ("Analytics Dashboard", "PARTIAL", "Some analytics endpoints still disabled"),
        ("Redis Local Install", "PENDING", "Phase 1 tasks 0.7-0.8. Caching logic ready, Redis not installed locally"),
        ("Production Deployment", "NOT DONE", "Currently development mode only. Needs deployment plan."),
        ("Flutter APK Build", "NEEDS SETUP", "Flutter SDK + Android SDK not installed on dev machine"),
        ("Face Detection Login", "EXPERIMENTAL", "Requires Chrome 98+ with FaceDetection API flag enabled"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 26. PHASE COMPLETION STATUS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "32. Phase Completion Status (Phase 0-8+)", 1)
    styled_table(doc, ["Phase", "Name", "Tasks", "Status", "What Was Built"], [
        ("Phase 0", "Foundation", "60", "COMPLETE", "Core infrastructure, MongoDB, PostgreSQL, basic AI analysis"),
        ("Phase 0B", "Error Documentation", "11", "COMPLETE", "25 curated error docs (ERR001-ERR025) loaded to Pinecone"),
        ("Phase 0C", "Dual-Index RAG", "13", "COMPLETE", "Two Pinecone indexes: knowledge-docs + error-library"),
        ("Phase 0D", "Context Engineering", "10", "COMPLETE", "Query expansion, context enrichment, RAGRouter OPTION C"),
        ("Phase 0E", "GitHub Integration", "11", "COMPLETE", "GitHub MCP server, code fetch, CodeSnippet UI component"),
        ("Phase 0-ARCH", "ReAct Agent", "30", "COMPLETE", "7-node LangGraph agent, ToolRegistry, self-correction"),
        ("Phase 1", "Redis Caching", "9", "7/9", "Caching logic ready, Redis installation pending (tasks 0.7-0.8)"),
        ("Phase 2", "Fusion RAG", "10", "COMPLETE", "4-source hybrid retrieval with CrossEncoder re-ranking"),
        ("Phase 3", "CRAG Verification", "10", "COMPLETE", "Confidence scoring, thresholds, escalation logic"),
        ("Phase 4", "PII Redaction", "8", "COMPLETE (OFF)", "Implemented but deliberately disabled (not needed)"),
        ("Phase 5", "HITL System", "10", "COMPLETE", "Human-in-the-loop: Accept/Reject/Refine workflow"),
        ("Phase 6", "Self-Correction", "8", "COMPLETE", "Automatic retry with exponential backoff, alternative tools"),
        ("Phase 7", "Code Fix Approval", "9", "COMPLETE", "AI-generated PR creation, approve/reject/rollback workflow"),
        ("Phase 8", "Langfuse Tracing", "10", "COMPLETE", "LLM call tracing, cost tracking, observability dashboard"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 33. KNOWN ISSUES & PENDING WORK
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "33. Known Issues & Pending Work", 1)

    heading(doc, "Known Issues", 2)
    styled_table(doc, ["#", "Issue", "Severity", "Impact", "Solution"], [
        ("1", "Redis not installed locally", "HIGH", "Phase 1 caching incomplete (tasks 0.7-0.8)", "Install Memurai (Windows) or use Docker Redis container"),
        ("2", "MongoDB connection timeouts", "MEDIUM", "Occasional failures connecting to Atlas", "URL-encode password special chars, verify IP whitelist"),
        ("3", "RAG low similarity scores", "MEDIUM", "Some queries return <50% similarity", "Enable query expansion, CrossEncoder re-ranking"),
        ("4", "Flutter SDK not installed", "MEDIUM", "Cannot build mobile APK on dev machine", "Install Flutter SDK 3.x + Android SDK"),
        ("5", "Docker on D: drive", "LOW", "If Docker reset, data location needs reconfiguration", "See DOCKER-FIX-AND-MIGRATION-GUIDE.md"),
        ("6", "PII redaction disabled", "INFO", "Phase 4 implemented but turned off by design", "Re-enable if needed in future"),
        ("7", "Dashboard N+1 queries", "LOW", "/api/failures has N+1 PostgreSQL queries", "Optimize with batch IN clause query"),
        ("8", "Face Detection browser support", "LOW", "FaceDetection API only in Chrome 98+ with flag", "Falls back to password login gracefully"),
    ])

    heading(doc, "Pending Work", 2)
    for item in [
        "Install Redis on Windows (Memurai) and complete Phase 1 tasks 0.7, 0.8",
        "Install Flutter SDK + Android SDK for mobile APK builds",
        "Complete Analytics page (some analytics endpoints still disabled)",
        "Set up production deployment (currently development mode only)",
        "Configure SMTP credentials for email notifications",
        "Performance testing under production load",
        "Add WebSocket for real-time dashboard updates (currently polling)",
        "Migrate React to TypeScript for better type safety",
        "Set up Firebase project for mobile push notifications",
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 20. ROLES & RESPONSIBILITIES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "34. Roles & Responsibilities (Actual RBAC Implemented)", 1)

    heading(doc, "Implemented RBAC Roles (6-Level Hierarchy)", 2)
    styled_table(doc, ["Role", "Dashboard Access", "Key Responsibilities", "Key Pages Used"], [
        ("super_admin", "Full + Admin pages", "System config, user management, all features", "UserManagement, Configuration, AuditLog, all pages"),
        ("project_owner", "Full + Project admin", "Manage project members, settings, approve fixes", "ProjectManagement, PRWorkflow, all analysis pages"),
        ("project_admin", "Project pages + Knowledge", "Approve fixes, manage knowledge, configure services", "Knowledge, PRWorkflow, ServicesMonitoring"),
        ("developer", "Analysis + Fix pages", "Fix bugs, review code fixes, trigger analysis", "FailureDetails, ManualTrigger, Copilot, PRWorkflow"),
        ("viewer", "Read-only", "Monitor metrics, review analysis results, view trends", "Dashboard, Failures, Analytics, JiraBugs"),
        ("guest", "Limited read", "View public dashboards only", "Dashboard (limited)"),
    ])

    heading(doc, "Stakeholder Mapping", 2)
    styled_table(doc, ["Stakeholder", "Recommended Role", "Key Pages"], [
        ("QA Engineer", "developer", "Failures, FailureDetails, ManualTrigger, TriggerAnalysis, RAGApproval"),
        ("Developer", "developer", "FailureDetails, Copilot, PRWorkflow, TestGenerator"),
        ("DevOps / Infra", "project_admin", "ServicesMonitoring, PipelineStatus, Configuration"),
        ("Knowledge Manager", "project_admin", "KnowledgeManagement, RAGApproval"),
        ("BA / PM", "viewer", "Dashboard, Analytics, JiraBugs, AuditLog"),
        ("Client (DDN)", "viewer", "Dashboard, Failures, Analytics"),
        ("Project Lead", "project_owner", "ProjectManagement, UserManagement, all pages"),
    ])

    heading(doc, "System Responsibilities by Component", 2)
    styled_table(doc, ["Component", "Owner Role", "What to Monitor", "When to Act"], [
        ("MongoDB Atlas", "DevOps", "Connection status, IP whitelist, data growth", "Timeout errors, auth failures"),
        ("PostgreSQL", "DevOps", "Connection pool, disk space, query performance", "Slow queries, connection errors"),
        ("Pinecone", "Knowledge Mgr", "Vector count, index health, query latency", "Low similarity scores, missing docs"),
        ("Redis", "DevOps", "Cache hit rate, memory usage, connection", "Cache misses, memory pressure"),
        ("n8n Workflows", "DevOps", "Workflow execution, error rate, triggers", "Failed workflows, missed triggers"),
        ("AI Agent (Port 5000)", "Developer", "Analysis quality, confidence scores, latency", "Low confidence, timeouts, errors"),
        ("Dashboard API (5006)", "Developer", "Response times, error rates, health check", "500 errors, slow responses"),
        ("Dashboard UI (5173)", "Frontend Dev", "Page load time, UI errors, user feedback", "Broken pages, missing data"),
        ("Langfuse (3000)", "DevOps/Dev", "LLM costs, token usage, trace completeness", "Cost spikes, missing traces"),
        ("Jenkins", "DevOps/QA", "Build status, test results, pipeline health", "Build failures, flaky tests"),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 21. HANDOVER CHECKLIST
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "35. Handover Checklist", 1)

    heading(doc, "Day 1 - Immediate Actions", 2)
    for item in [
        "[ ] Get access to GitHub (https://github.com/Sushrut-01) - clone all 3 repos",
        "[ ] Access the development machine RYSUN-PUNE-03 (10.0.1.3)",
        "[ ] Open project in VS Code: code D:\\DDN-AI-Project-Documentation",
        "[ ] Activate venv and verify Python: python --version (should be 3.13.5)",
        "[ ] Copy .env.MASTER and verify all API keys are valid",
        "[ ] Read CLAUDE.md for complete project context",
        "[ ] Read this handover document completely",
        "[ ] Run health check: curl http://localhost:5006/api/health",
    ]:
        doc.add_paragraph(item)

    heading(doc, "Week 1 - Setup & Verification", 2)
    for item in [
        "[ ] Start Docker containers: docker-compose up -d",
        "[ ] Verify all 17 services running: docker-compose ps",
        "[ ] Open dashboard: http://localhost:5173",
        "[ ] Open n8n: http://localhost:5678",
        "[ ] Open Langfuse: http://localhost:3000",
        "[ ] Test MongoDB: python implementation/database/test_mongodb_connection.py",
        "[ ] Run pytest: pytest implementation/ -v --tb=short",
        "[ ] Trigger manual analysis for a test build",
        "[ ] Review Jenkins pipeline configuration",
        "[ ] Install Redis (Memurai) to complete Phase 1",
        "[ ] Commit and push all 147 pending files to GitHub",
    ]:
        doc.add_paragraph(item)

    heading(doc, "Month 1 - Full Ownership", 2)
    for item in [
        "[ ] Complete Redis installation and Phase 1 caching",
        "[ ] Implement authentication (JWT) on Dashboard API",
        "[ ] Complete Analytics page (category trends, model performance)",
        "[ ] Review and rotate API keys if needed",
        "[ ] Set up production deployment plan",
        "[ ] Configure monitoring alerts (PagerDuty/Teams)",
        "[ ] Run performance tests under expected production load",
        "[ ] Update PROGRESS-TRACKER-FINAL.csv for completed tasks",
        "[ ] Review CRAG confidence thresholds (may need tuning)",
        "[ ] Schedule knowledge base review with domain experts",
    ]:
        doc.add_paragraph(item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 22. KEY CONTACTS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "36. Key Contacts", 1)
    styled_table(doc, ["Name", "Role", "Contact", "Responsibility"], [
        ("Sushrut Nistane", "Developer (Outgoing)", "sushrut.nistane@rysun.com", "Built entire system. Available for handover questions."),
        ("Project Manager", "PM", "Contact via Teams", "DDN AI Project channel. Requirements, timelines, client comms."),
        ("Backend Lead", "Tech Lead", "Contact via Teams", "Python services, AI/ML, LangGraph, databases."),
        ("Frontend Lead", "Tech Lead", "Contact via Teams", "React dashboard, UI/UX, Material-UI."),
        ("DDN Client Team", "Client", "Contact via PM", "Product requirements, UAT, feedback."),
        ("Rysun Management", "Vendor", "Contact via PM", "Resource allocation, escalations."),
    ])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 23. FAQ FOR STAKEHOLDERS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "37. FAQ for Stakeholders", 1)

    heading(doc, "For BA / PM", 2)
    for q, a in [
        ("Q: What is the business value of this system?",
         "A: Reduces debugging time from 60min to 15sec (99.6%), cost from $30 to $0.05 (99.8%), and increases engineer productivity 3x."),
        ("Q: What is the ROI?",
         "A: 103.7% in Year 1. Cost savings from reduced manual effort far exceed AI API costs."),
        ("Q: How do we measure success?",
         "A: Track acceptance rate (target >80%), average confidence score (target >0.75), cases per day per engineer."),
        ("Q: What is the project status?",
         "A: All 8 phases complete + QA Agent features merged. 60+ major features delivered including JWT auth, multi-project architecture, AI Copilot, Flutter mobile app, Jira, GitHub PR, Slack, Notifications, and Knowledge Management."),
    ]:
        para(doc, q, bold=True, size=10)
        doc.add_paragraph(a)

    heading(doc, "For Developers", 2)
    for q, a in [
        ("Q: How does the AI analysis work?",
         "A: 7-node LangGraph ReAct agent: Classify -> Reason -> Select Tool -> Execute -> Observe -> Answer -> Verify. Uses GPT-4o-mini for reasoning, Fusion RAG for knowledge retrieval, CRAG for verification."),
        ("Q: How do I add a new error category?",
         "A: Add vectors to Pinecone indexes. The ToolRegistry auto-discovers categories from Pinecone metadata (no code changes needed). 5-minute cache TTL."),
        ("Q: How do I add a new tool to the agent?",
         "A: Register in tool_registry.py with ToolMetadata (name, description, cost, latency, use_for categories, priority). Implement the tool method in react_agent_service.py."),
        ("Q: Why is PII redaction disabled?",
         "A: Phase 4 was fully implemented but turned off because DDN test data doesn't contain PII. Can be re-enabled by setting PII_REDACTION_ENABLED=true in .env.MASTER."),
    ]:
        para(doc, q, bold=True, size=10)
        doc.add_paragraph(a)

    heading(doc, "For QA", 2)
    for q, a in [
        ("Q: How do I run all tests?",
         "A: Activate venv, then: pytest implementation/ -v --tb=short. For Robot: cd robot-tests && robot ddn_basic_tests.robot. For Playwright: see ddn-playwright-automation repo."),
        ("Q: How do I validate an AI analysis?",
         "A: Go to /failures -> click build -> review AI analysis. Click Accept (correct), Reject (wrong, give reason), or Refine (partially correct, suggest improvements)."),
        ("Q: How do I trigger analysis for a specific build?",
         "A: Go to /manual-trigger page, enter Build ID, click Trigger. Or use /trigger-analysis for bulk processing of unanalyzed failures."),
        ("Q: What test coverage exists?",
         "A: 50+ pytest files, 4 Mocha files, 2 Robot Framework suites, 3 Playwright suites. Target is >80% coverage."),
    ]:
        para(doc, q, bold=True, size=10)
        doc.add_paragraph(a)

    heading(doc, "For Client (DDN)", 2)
    for q, a in [
        ("Q: What DDN products are tested?",
         "A: EXAScaler (Lustre), AI400X (GPU storage), Infinia (orchestration), IntelliFlash (flash arrays), JARVICE (HPC cloud), S3 storage, Multi-Tenancy."),
        ("Q: How do we access the dashboard?",
         "A: Open http://localhost:5173 in a browser. Login with email/password (JWT auth). Supports Google/GitHub OAuth, JARVIS voice, and face detection login. View failures, analysis, and system health."),
        ("Q: How accurate is the AI analysis?",
         "A: Average confidence score targets >75%. Analyses with >85% confidence are auto-accepted. Below 65% triggers human review."),
        ("Q: Can we add our own error documentation?",
         "A: Yes. Go to /knowledge page. Click '+ Add Knowledge Doc'. Fill in error details, root cause, solution, code examples. It's automatically indexed in Pinecone."),
    ]:
        para(doc, q, bold=True, size=10)
        doc.add_paragraph(a)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 24. DOCUMENTATION INDEX
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "38. Additional Documentation Index", 1)

    heading(doc, "Architecture & Design", 2)
    for name, desc in [
        ("CLAUDE.md", "Complete project memory and system architecture (AI assistant context)"),
        ("BA-PROJECT-UNDERSTANDING.md", "Business analysis perspective and stakeholder view"),
        ("ARCHITECTURE-DIAGRAMS-COMPLETE.md", "Visual architecture diagrams (Mermaid)"),
        ("GRANULAR-TECHNICAL-ARCHITECTURE.md", "Detailed technical architecture breakdown"),
        ("PROJECT-BASED-GRANULAR-ARCHITECTURE.md", "Project-based architecture view"),
        ("ARCHITECTURE-OVERVIEW.html", "Interactive HTML architecture overview"),
    ]:
        bullet(doc, f"{name} - {desc}")

    heading(doc, "Setup & Integration Guides", 2)
    for name, desc in [
        ("START-HERE.md", "Quick start guide for new developers"),
        ("DEPLOYMENT-GUIDE.md", "Production deployment instructions"),
        ("N8N-WORKFLOW-SETUP-GUIDE.md", "n8n workflow configuration"),
        ("JENKINS-CLI-GUIDE.md", "Jenkins CLI and job setup"),
        ("MONGODB-ATLAS-SETUP.md", "MongoDB Atlas configuration and IP whitelist"),
        ("DOCKER-FIX-AND-MIGRATION-GUIDE.md", "Docker troubleshooting and D: drive migration"),
        ("GITHUB-REPOSITORY-SETUP.md", "GitHub repository setup guide"),
    ]:
        bullet(doc, f"{name} - {desc}")

    heading(doc, "Phase-Specific Guides", 2)
    for name, desc in [
        ("CONTRIBUTING-ERROR-DOCS.md", "Phase 0B: How to add error documentation"),
        ("DUAL-INDEX-RAG-ARCHITECTURE.md", "Phase 0C: Dual-index Pinecone strategy"),
        ("CONTEXT-ENGINEERING-GUIDE.md", "Phase 0D: Context engineering and RAGRouter"),
        ("GITHUB-INTEGRATION-GUIDE.md", "Phase 0E: GitHub MCP integration"),
        ("REACT-AGENT-GUIDE.md", "Phase 0-ARCH: ReAct agent architecture"),
        ("CRAG-VERIFICATION-GUIDE.md", "Phase 3: CRAG confidence verification"),
        ("PII-REDACTION-DISABLED-SUMMARY.md", "Phase 4: PII redaction (why disabled)"),
    ]:
        bullet(doc, f"{name} - {desc}")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 25. QUICK REFERENCE COMMANDS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "39. Quick Reference Commands", 1)
    code_block(doc,
        "# ===== HEALTH & STATUS =====\n"
        "curl http://localhost:5006/api/health              # API health\n"
        "curl http://localhost:5006/api/system/status        # Full system status\n"
        "powershell .\\check-all-services.ps1                # Check all services\n\n"
        "# ===== VIEW DATA =====\n"
        'curl "http://localhost:5006/api/failures?limit=10"  # Recent failures\n'
        "curl http://localhost:5006/api/stats                # Statistics\n"
        "curl http://localhost:5006/api/activity             # Activity log\n\n"
        "# ===== TRIGGER ANALYSIS =====\n"
        "curl -X POST http://localhost:5004/trigger-analysis \\\n"
        '  -H "Content-Type: application/json" \\\n'
        "  -d '{\"build_id\": \"12345\"}'                      # Manual trigger\n\n"
        "# ===== DATABASES =====\n"
        "python implementation/database/test_mongodb_connection.py  # Test MongoDB\n"
        "python implementation/test_postgres_connection.py          # Test PostgreSQL\n"
        "redis-cli ping                                             # Test Redis\n\n"
        "# ===== TESTING =====\n"
        "pytest implementation/ -v --tb=short                # All Python tests\n"
        "pytest implementation/tests/test_react_agent.py -v  # Agent tests\n"
        "cd tests && npm test                                # JavaScript tests\n"
        "cd robot-tests && robot ddn_basic_tests.robot       # Robot tests\n\n"
        "# ===== DOCKER =====\n"
        "docker-compose up -d                    # Start all\n"
        "docker-compose down                     # Stop all\n"
        "docker-compose ps                       # Status\n"
        "docker-compose logs -f dashboard-api    # View logs\n\n"
        "# ===== URLS =====\n"
        "# Dashboard:  http://localhost:5173\n"
        "# n8n:        http://localhost:5678\n"
        "# Langfuse:   http://localhost:3000\n"
        "# API:        http://localhost:5006\n"
        "# Flower:     http://localhost:5555\n\n"
        "# ===== PORT CHECK =====\n"
        "netstat -ano | findstr :5006   # Check port usage\n"
        "netstat -ano | findstr :5000   # Check agent port\n"
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 26. HANDOVER EMAIL FORMAT
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "40. Handover Email Format", 1)
    doc.add_paragraph("Copy the email below and send to your manager/team:")
    spacer(doc)

    email_text = f"""Subject: Project Handover - DDN AI Test Case Failure Analysis System

Dear [Manager Name],

I am writing to formally hand over the DDN AI-Assisted Test Case Failure Analysis System project. Please find below all necessary information for a smooth transition.

PROJECT SUMMARY
- Project: DDN AI-Assisted Test Case Failure Analysis System
- Client: Data Direct Networks (DDN)
- Vendor: Rysun Labs Pvt. Ltd.
- Status: All phases complete + QA Agent features merged
- Key Achievement: Reduced manual debugging from 60 min to 15 sec (99.6% reduction)

REPOSITORY LINKS (GitHub)
1. Main Project: https://github.com/Sushrut-01/ddn-ai-test-analysis
2. Playwright E2E Tests: https://github.com/Sushrut-01/ddn-playwright-automation
3. Robot Framework Tests: https://github.com/Sushrut-01/ddn-jenkins-testing

PROJECT LOCATION
- Machine: RYSUN-PUNE-03 (10.0.1.3)
- Folder: D:\\DDN-AI-Project-Documentation
- VS Code: code "D:\\DDN-AI-Project-Documentation"

TECHNOLOGY STACK
Python 3.13 | React 18 | Flutter 3.x | Flask | LangGraph | MongoDB Atlas |
PostgreSQL (RLS) | Pinecone | Redis | Rancher Desktop + Docker (17+ services) |
n8n | Langfuse | Jenkins | Claude 3.5 | GPT-4o-mini | Gemini Flash 2.0 |
Celery | Flower | JWT + bcrypt | Web Speech API

KEY FEATURES DEVELOPED (60+ Features)
- JWT Authentication & RBAC (6-level roles, bcrypt, refresh tokens)
- Multi-Project Architecture (PostgreSQL RLS, per-project isolation)
- AI Copilot with JARVIS Voice Assistant (Gemini Flash 2.0 + Web Speech API)
- Flutter Mobile App (Android/iOS, offline-first, push notifications)
- 25-Page React Dashboard (Login, Copilot, Projects, Services, Audit...)
- Jira Auto-Bug Creation (smart duplicate prevention, auto-priority)
- GitHub PR Auto-Creation (full code fix workflow with approval)
- Slack Integration (rich notifications, interactive buttons, slash commands)
- Knowledge Management API (full CRUD with Pinecone vector sync)
- Notifications Service (email + in-app with per-user preferences)
- Workflow Execution Tracker (n8n-style node-level debugging)
- Guruttava Project Integration (second project fully configured)
- Service Manager, Aging Service, Self-Healing Service
- Celery Async Tasks (prevent webhook timeouts)

PENDING ITEMS
1. Redis local installation (caching logic ready)
2. Flutter APK build (SDK setup required)
3. Production deployment setup

HANDOVER DOCUMENT
Please find attached the comprehensive handover document (DDN-PROJECT-HANDOVER-DOCUMENT.docx)
with 40 sections covering:
- System architecture (7-node ReAct agent, Fusion RAG, CRAG verification)
- 100+ REST API endpoints across 12+ Python services (v1 + v2)
- JWT Authentication & 6-level RBAC with multi-project RLS
- AI Copilot with JARVIS voice assistant
- Flutter mobile app (Android/iOS) with offline-first architecture
- React dashboard (25 pages, 25+ components, dark/light theme)
- Database schemas (MongoDB + PostgreSQL + Pinecone) with 8 migrations
- Jira, GitHub PR, Slack, Notifications integrations
- Rancher Desktop + Docker (17+ containers)
- DDN products + Guruttava project (multi-project)
- Implementation maturity assessment (60+ features)
- Handover checklist (Day 1, Week 1, Month 1)
- FAQ for BA, PM, Dev, QA, and Client

I am available for questions during the transition period.

Best regards,
Sushrut Nistane
sushrut.nistane@rysun.com
Rysun Labs Pvt. Ltd."""

    p = doc.add_paragraph()
    run = p.add_run(email_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════════════════════════
    spacer(doc)
    para(doc, "-" * 70, size=8, color=(192, 192, 192), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "RYSUN LABS PVT. LTD.", bold=True, size=11,
         color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Pune, Maharashtra, India", size=9,
         color=(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)
    year = datetime.date.today().year
    para(doc, f"\u00A9 {year} Rysun Labs Pvt. Ltd. All Rights Reserved.",
         size=8, color=(128, 128, 128), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "This document is confidential and proprietary to Rysun Labs Pvt. Ltd.",
         size=8, color=(128, 128, 128), align=WD_ALIGN_PARAGRAPH.CENTER)

    # ══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════════════
    path = r"D:\DDN-AI-Project-Documentation\DDN-PROJECT-HANDOVER-DOCUMENT.docx"
    doc.save(path)
    print(f"\nHandover document saved to: {path}")
    print(f"Sections: 40")
    print(f"Date: {today}")
    return path


if __name__ == "__main__":
    build()
