"""
Extract content from Word document and convert to markdown
"""
from docx import Document
import sys

def extract_docx_to_markdown(docx_path, md_path):
    """Extract text from docx and save as markdown"""
    doc = Document(docx_path)

    with open(md_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                f.write('\n')
                continue

            # Check paragraph style for headings
            style = para.style.name.lower()
            if 'heading 1' in style:
                f.write(f'# {text}\n\n')
            elif 'heading 2' in style:
                f.write(f'## {text}\n\n')
            elif 'heading 3' in style:
                f.write(f'### {text}\n\n')
            elif 'heading 4' in style:
                f.write(f'#### {text}\n\n')
            elif 'heading 5' in style:
                f.write(f'##### {text}\n\n')
            elif 'heading 6' in style:
                f.write(f'###### {text}\n\n')
            else:
                f.write(f'{text}\n\n')

        # Extract tables
        for table in doc.tables:
            f.write('\n')
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                f.write('| ' + ' | '.join(cells) + ' |\n')
                if i == 0:
                    f.write('| ' + ' | '.join(['---'] * len(cells)) + ' |\n')
            f.write('\n')

    print(f"Successfully extracted content to: {md_path}")

if __name__ == '__main__':
    docx_file = r'C:\DDN-AI-Project-Documentation\implementation\DDN_Document_Gap_Analysis_And_Integration_Plan.docx'
    md_file = r'C:\DDN-AI-Project-Documentation\DDN_Document_Gap_Analysis_And_Integration_Plan.md'

    extract_docx_to_markdown(docx_file, md_file)
